"""
AI-powered semantic search using INFINIA accelerated by NVIDIA GPU technology
Supports: Images, Videos, Documents with NVIDIA's GPU computing platform
"""

import gradio as gr
import boto3
from botocore.client import Config
from botocore.exceptions import ClientError
import io
import json
import tempfile
import os
from PIL import Image
import cv2
import numpy as np
from datetime import datetime
import hashlib
import warnings
from typing import List, Dict, Tuple, Optional, Any
import threading
import time
from pathlib import Path
import mimetypes
import unicodedata
import re

warnings.filterwarnings('ignore')

# AI/ML Libraries
try:
    from transformers import pipeline, BlipProcessor, BlipForConditionalGeneration, AutoTokenizer, AutoModel, CLIPProcessor, CLIPModel
    import torch
    import torch.nn.functional as F
    TRANSFORMERS_AVAILABLE = True
except ImportError:
    TRANSFORMERS_AVAILABLE = False
    print("⚠️  Install transformers: pip install transformers torch")

# Document processing
try:
    import PyPDF2
    import docx
    DOCUMENT_PROCESSING_AVAILABLE = True
except ImportError:
    DOCUMENT_PROCESSING_AVAILABLE = False
    print("⚠️  Install document libs: pip install PyPDF2 python-docx")

# NVIDIA integrations
try:
    import nemo
    import nemo.collections.nlp as nemo_nlp
    NVIDIA_NEMO_AVAILABLE = True
except ImportError:
    NVIDIA_NEMO_AVAILABLE = False
    print("ℹ️  NVIDIA NeMo not available (optional)")

# Exa Filesystem
try:
    from exa_filesystem_manager import ExaFilesystemManager
    EXA_FILESYSTEM_AVAILABLE = True
except ImportError:
    EXA_FILESYSTEM_AVAILABLE = False
    print("⚠️  Exa Filesystem support not available. Install paramiko: pip install paramiko")

# Global variables
s3_manager = None
exa_manager = None
storage_backend = "s3"  # "s3" or "exa"
ai_analyzer = None
video_analyzer = None
document_analyzer = None


def sanitize_for_s3_metadata(text: str) -> str:
    """
    Sanitize text to be S3 metadata compliant (ASCII only).
    Converts Unicode characters to ASCII equivalents.
    """
    if not text:
        return ""
    
    # Replace common Unicode characters with ASCII equivalents
    replacements = {
        # Smart quotes
        '\u201c': '"',  # Left double quotation mark
        '\u201d': '"',  # Right double quotation mark
        '\u2018': "'",  # Left single quotation mark
        '\u2019': "'",  # Right single quotation mark
        # Dashes
        '\u2014': '-',  # Em dash
        '\u2013': '-',  # En dash
        '\u2012': '-',  # Figure dash
        # Ellipsis
        '\u2026': '...',  # Horizontal ellipsis
        # Spaces
        '\u00a0': ' ',  # Non-breaking space
        '\u2009': ' ',  # Thin space
        # Bullets
        '\u2022': '*',  # Bullet
        '\u2023': '*',  # Triangular bullet
        # Other common symbols
        '\u00ae': '(R)',  # Registered trademark
        '\u2122': '(TM)',  # Trademark
        '\u00a9': '(C)',  # Copyright
    }
    
    # Apply replacements
    for unicode_char, ascii_char in replacements.items():
        text = text.replace(unicode_char, ascii_char)
    
    # Normalize Unicode to decomposed form, then remove combining characters
    text = unicodedata.normalize('NFKD', text)
    text = text.encode('ascii', 'ignore').decode('ascii')
    
    # Remove any remaining control characters except newlines and tabs
    text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    
    # Truncate if too long (S3 metadata has limits)
    max_length = 2000  # Conservative limit
    if len(text) > max_length:
        text = text[:max_length - 3] + '...'
    
    return text.strip()


class InfiniaS3Manager:
    def __init__(self):
        self.client = None
        self.bucket = None
        self.connection_status = "Not Connected"
        self.endpoint = None
    
    def connect(self, endpoint: str, access_key: str, secret_key: str, 
                bucket: str, region: str = "us-east-1") -> Tuple[bool, str]:
        """Establish connection to DDN Infinia Object Storage"""
        try:
            self.endpoint = endpoint
            self.client = boto3.client(
                's3',
                endpoint_url=endpoint,
                aws_access_key_id=access_key,
                aws_secret_access_key=secret_key,
                region_name=region,
                config=Config(signature_version='s3v4'),
                verify=False
            )
            
            # Test connection
            self.client.head_bucket(Bucket=bucket)
            self.bucket = bucket
            self.connection_status = "Connected"
            
            # Create bucket if doesn't exist
            try:
                self.client.create_bucket(Bucket=bucket)
            except ClientError:
                pass
            
            return True, f"✅ Successfully connected to Infinia bucket: {bucket}"
        except Exception as e:
            self.connection_status = "Connection Failed"
            return False, f"❌ Connection failed: {str(e)}"
    
    def upload_file(self, file_path: str, object_key: str, metadata: Dict, 
                    content_type: str = None) -> Tuple[bool, str]:
        """Upload file with comprehensive metadata to Infinia"""
        try:
            extra_args = {'Metadata': metadata}
            if content_type:
                extra_args['ContentType'] = content_type
            
            with open(file_path, 'rb') as f:
                self.client.put_object(
                    Bucket=self.bucket,
                    Key=object_key,
                    Body=f,
                    **extra_args
                )
            return True, f"Uploaded: {object_key}"
        except Exception as e:
            return False, f"Upload failed: {str(e)}"
    
    def search_objects(self, query: str, modality_filter: str = "all") -> List[Dict]:
        """Search objects based on metadata with modality filtering and keyword extraction"""
        try:
            results = []
            response = self.client.list_objects_v2(Bucket=self.bucket)
            
            if 'Contents' not in response:
                return results
            
            # Auto-detect modality from query if filter is "all"
            query_lower = query.lower()
            if modality_filter == "all":
                if "image" in query_lower or "images" in query_lower or "picture" in query_lower or "photo" in query_lower:
                    modality_filter = "image"
                elif "video" in query_lower or "videos" in query_lower:
                    modality_filter = "video"
                elif "document" in query_lower or "documents" in query_lower or "doc" in query_lower or "pdf" in query_lower:
                    modality_filter = "document"
            
            # Extract keywords from query (remove common words)
            stop_words = {
                'show', 'me', 'the', 'where', 'is', 'are', 'in', 'a', 'an', 'and', 'or', 'of', 'to', 'with', 
                'images', 'videos', 'documents', 'image', 'video', 'document', 'all', 'find', 'search',
                'picture', 'pictures', 'photo', 'photos', 'file', 'files'
            }
            query_words = [word.lower() for word in query.split() if word.lower() not in stop_words and len(word) > 2]
            
            if not query_words:
                # If no keywords, use full query
                query_words = [query_lower]
            
            for obj in response['Contents']:
                try:
                    head_response = self.client.head_object(
                        Bucket=self.bucket,
                        Key=obj['Key']
                    )
                    
                    metadata = head_response.get('Metadata', {})
                    
                    # Filter by modality
                    if modality_filter != "all":
                        obj_modality = metadata.get('modality', 'unknown')
                        if obj_modality != modality_filter:
                            continue
                    
                    # Search for keywords in metadata values
                    match_score = 0
                    matched_fields = []
                    
                    # Only search in relevant metadata fields (not in descriptions or long text)
                    relevant_fields = ['caption', 'ai_caption', 'scene_1', 'scene_2', 'scene_3', 
                                      'dominant_scene', 'detected_objects', 'summary', 'ai_summary', 'custom_tags']
                    
                    searchable_text = ''
                    for field in relevant_fields:
                        if field in metadata:
                            searchable_text += ' ' + str(metadata[field]).lower()
                    
                    # Check each keyword
                    for keyword in query_words:
                        if keyword in searchable_text:
                            match_score += 1
                            # Find which field matched
                            for key in relevant_fields:
                                if key in metadata:
                                    field_value = str(metadata[key]).lower()
                                    
                                    # Special handling for custom_tags - check individual tags
                                    if key == 'custom_tags':
                                        # Split by comma and check each tag individually
                                        individual_tags = [t.strip() for t in field_value.split(',')]
                                        if keyword in individual_tags or keyword in field_value:
                                            matched_fields.append(f"{key}: {metadata[key]}")
                                            break
                                    elif keyword in field_value:
                                        matched_fields.append(f"{key}: {metadata[key]}")
                                        break
                    
                    if match_score > 0:
                        results.append({
                            'key': obj['Key'],
                            'size': obj['Size'],
                            'last_modified': obj['LastModified'],
                            'metadata': metadata,
                            'match_score': match_score,
                            'matched_fields': list(set(matched_fields))[:5]  # Remove duplicates, limit to 5
                        })
                
                except Exception as e:
                    continue
            
            # Sort by match score
            results.sort(key=lambda x: x['match_score'], reverse=True)
            return results
            
        except Exception as e:
            print(f"Search error: {str(e)}")
            return []
    
    def get_object(self, object_key: str) -> Optional[bytes]:
        """Retrieve object data from Infinia"""
        try:
            response = self.client.get_object(Bucket=self.bucket, Key=object_key)
            return response['Body'].read()
        except Exception as e:
            print(f"Get object error: {e}")
            return None
    
    def get_object_metadata(self, object_key: str) -> Optional[Dict]:
        """Get object metadata"""
        try:
            response = self.client.head_object(Bucket=self.bucket, Key=object_key)
            return response.get('Metadata', {})
        except Exception as e:
            print(f"Get metadata error: {e}")
            return None
    
    def get_all_objects(self, modality_filter: str = "all") -> List[Dict]:
        """Get all objects with optional modality filtering"""
        try:
            objects = []
            response = self.client.list_objects_v2(Bucket=self.bucket)
            
            if 'Contents' not in response:
                return objects
            
            for obj in response['Contents']:
                try:
                    head_response = self.client.head_object(
                        Bucket=self.bucket,
                        Key=obj['Key']
                    )
                    metadata = head_response.get('Metadata', {})
                    
                    # Filter by modality if specified
                    if modality_filter != "all":
                        if metadata.get('modality', 'unknown') != modality_filter:
                            continue
                    
                    objects.append({
                        'key': obj['Key'],
                        'size': obj['Size'],
                        'last_modified': obj['LastModified'],
                        'modality': metadata.get('modality', 'unknown')
                    })
                except:
                    continue
            
            return objects
        except Exception as e:
            print(f"List objects error: {e}")
            return []


class ImageAnalyzer:
    """Advanced image analysis using NVIDIA and Hugging Face models"""
    
    def __init__(self):
        self.models_loaded = False
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        
        if not TRANSFORMERS_AVAILABLE:
            print("⚠️  Transformers not available. Image analysis disabled.")
            return
        
        try:
            print("Loading BLIP for captioning...")
            self.caption_processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base")
            self.caption_model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base").to(self.device)
            
            print("Loading scene classifier...")
            self.scene_classifier = pipeline("image-classification", model="google/vit-base-patch16-224", device=0 if self.device == "cuda" else -1)
            
            # Load CLIP for semantic search
            print("Loading CLIP for semantic search...")
            self.clip_processor = CLIPProcessor.from_pretrained("openai/clip-vit-base-patch32")
            self.clip_model = CLIPModel.from_pretrained("openai/clip-vit-base-patch32").to(self.device)
            
            self.models_loaded = True
            print("✅ All image models loaded successfully")
            
        except Exception as e:
            print(f"⚠️  Model loading error: {e}")
            self.models_loaded = False
    
    def generate_caption(self, image: Image.Image) -> str:
        """Generate image caption using BLIP"""
        if not self.models_loaded:
            return "AI models not loaded"
        
        try:
            inputs = self.caption_processor(image, return_tensors="pt").to(self.device)
            outputs = self.caption_model.generate(**inputs, max_new_tokens=50)
            caption = self.caption_processor.decode(outputs[0], skip_special_tokens=True)
            return caption
        except Exception as e:
            return f"Caption error: {e}"
    
    def classify_scene(self, image: Image.Image, top_k: int = 5) -> List[Dict]:
        """Classify image scene"""
        if not self.models_loaded:
            return []
        
        try:
            results = self.scene_classifier(image, top_k=top_k)
            return results
        except Exception as e:
            print(f"Scene classification error: {e}")
            return []
    
    def detect_objects(self, image: Image.Image) -> str:
        """Detect objects in image using scene classification"""
        scene_results = self.classify_scene(image, top_k=3)
        if scene_results:
            objects = [r['label'] for r in scene_results]
            return ', '.join(objects)
        return 'unknown'
    
    def compute_clip_similarity(self, image: Image.Image, text: str) -> float:
            """Compute CLIP similarity between image and text query using cosine similarity"""
            if not self.models_loaded:
                return 0.0
            
            try:
                inputs = self.clip_processor(text=[text], images=image, return_tensors="pt", padding=True).to(self.device)
                
                with torch.no_grad():
                    outputs = self.clip_model(**inputs)
                    
                    # Get image and text embeddings
                    image_embeds = outputs.image_embeds
                    text_embeds = outputs.text_embeds
                    
                    # Normalize embeddings
                    image_embeds = F.normalize(image_embeds, p=2, dim=-1)
                    text_embeds = F.normalize(text_embeds, p=2, dim=-1)
                    
                    # Compute cosine similarity (dot product of normalized vectors)
                    similarity = (image_embeds @ text_embeds.T).squeeze().item()
                    
                    # Convert from [-1, 1] range to [0, 1] range for consistency
                    similarity = (similarity + 1) / 2
                
                return similarity
            except Exception as e:
                print(f"CLIP similarity error: {e}")
                return 0.0
    
    def generate_clip_embedding(self, image: Image.Image) -> np.ndarray:
        """Generate CLIP image embedding vector for semantic search
        
        Returns normalized 512-dimensional embedding vector
        """
        if not self.models_loaded:
            return np.array([])
        
        try:
            # Process image through CLIP
            inputs = self.clip_processor(images=image, return_tensors="pt").to(self.device)
            
            # Generate embedding without gradients (inference only)
            with torch.no_grad():
                image_features = self.clip_model.get_image_features(**inputs)
                # Normalize embedding for cosine similarity
                embedding = F.normalize(image_features, p=2, dim=1)
            
            # Return as numpy array (flatten to 1D)
            return embedding.cpu().numpy().flatten()
        
        except Exception as e:
            print(f"❌ CLIP image embedding error: {e}")
            return np.array([])
    
    def generate_clip_text_embedding(self, text: str) -> np.ndarray:
        """Generate CLIP text embedding vector for query matching
        
        Args:
            text: Search query text
            
        Returns normalized 512-dimensional embedding vector
        """
        if not self.models_loaded:
            return np.array([])
        
        try:
            # Process text through CLIP
            inputs = self.clip_processor(
                text=[text], 
                return_tensors="pt", 
                padding=True
            ).to(self.device)
            
            # Generate embedding without gradients (inference only)
            with torch.no_grad():
                text_features = self.clip_model.get_text_features(**inputs)
                # Normalize embedding for cosine similarity
                embedding = F.normalize(text_features, p=2, dim=1)
            
            # Return as numpy array (flatten to 1D)
            return embedding.cpu().numpy().flatten()
        
        except Exception as e:
            print(f"❌ CLIP text embedding error: {e}")
            return np.array([])
    
    def analyze_image(self, image: Image.Image) -> Dict[str, Any]:
        """Comprehensive image analysis"""
        try:
            width, height = image.size
            
            # AI analysis
            caption = self.generate_caption(image)
            scene_classification = self.classify_scene(image)
            objects = self.detect_objects(image)
            
            return {
                'caption': caption,
                'width': width,
                'height': height,
                'aspect_ratio': round(width / height, 2),
                'scene_classification': scene_classification,
                'detected_objects': objects,
                'dominant_scene': scene_classification[0]['label'] if scene_classification else 'unknown'
            }
        
        except Exception as e:
            print(f"Image analysis error: {e}")
            return {
                'caption': 'Analysis error',
                'error': str(e)
            }


class VideoAnalyzer:
    """Enhanced video analysis with CLIP-based semantic search and scene detection"""
    
    def __init__(self):
        if TRANSFORMERS_AVAILABLE:
            self.device = "cuda" if torch.cuda.is_available() else "cpu"
        else:
            self.device = "cpu"
        self.image_analyzer = None
    
    def set_image_analyzer(self, analyzer: ImageAnalyzer):
        """Set image analyzer for frame analysis"""
        self.image_analyzer = analyzer
    
    def extract_keyframes(self, video_path: str, num_frames: int = 10) -> List[Tuple[int, np.ndarray]]:
        """Extract key frames from video"""
        try:
            cap = cv2.VideoCapture(video_path)
            total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            fps = cap.get(cv2.CAP_PROP_FPS)
            
            if total_frames == 0:
                return []
            
            # Calculate frame indices to extract
            frame_indices = np.linspace(0, total_frames - 1, num_frames, dtype=int)
            
            keyframes = []
            for idx in frame_indices:
                cap.set(cv2.CAP_PROP_POS_FRAMES, idx)
                ret, frame = cap.read()
                if ret:
                    keyframes.append((idx, frame))
            
            cap.release()
            return keyframes
        
        except Exception as e:
            print(f"Keyframe extraction error: {e}")
            return []
    
    def analyze_video(self, video_path: str) -> Dict[str, Any]:
        """Comprehensive video analysis"""
        try:
            cap = cv2.VideoCapture(video_path)
            
            # Video properties
            fps = cap.get(cv2.CAP_PROP_FPS)
            frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
            height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
            duration = frame_count / fps if fps > 0 else 0
            
            cap.release()
            
            # Extract and analyze keyframes
            keyframes = self.extract_keyframes(video_path, num_frames=8)
            
            frame_analyses = []
            scene_labels = []
            
            if self.image_analyzer and self.image_analyzer.models_loaded:
                for frame_idx, frame in keyframes:
                    # Convert BGR to RGB
                    frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                    pil_image = Image.fromarray(frame_rgb)
                    
                    # Analyze frame
                    caption = self.image_analyzer.generate_caption(pil_image)
                    scene_class = self.image_analyzer.classify_scene(pil_image)
                    
                    timestamp = frame_idx / fps if fps > 0 else 0
                    
                    frame_analyses.append({
                        'frame_index': int(frame_idx),
                        'timestamp': round(timestamp, 2),
                        'caption': caption,
                        'scene': scene_class[0]['label'] if scene_class else 'unknown'
                    })
                    
                    if scene_class:
                        scene_labels.append(scene_class[0]['label'])
            
            # Generate video summary
            summary = self.generate_video_summary(frame_analyses)
            
            # Dominant scenes
            from collections import Counter
            scene_counts = Counter(scene_labels)
            dominant_scenes = [
                {'scene': scene, 'count': count}
                for scene, count in scene_counts.most_common(3)
            ]
            
            return {
                'duration': round(duration, 2),
                'fps': round(fps, 2),
                'frame_count': frame_count,
                'width': width,
                'height': height,
                'resolution': f"{width}x{height}",
                'keyframes_analyzed': len(frame_analyses),
                'frame_analyses': frame_analyses,
                'dominant_scenes': dominant_scenes,
                'video_summary': summary
            }
        
        except Exception as e:
            print(f"Video analysis error: {e}")
            return {
                'error': str(e),
                'video_summary': 'Analysis failed'
            }
    
    def generate_video_summary(self, frame_analyses: List[Dict]) -> str:
        """Generate natural language summary of video content"""
        if not frame_analyses:
            return "No frame analysis available"
        
        try:
            # Extract scenes and captions
            scenes = [fa['scene'] for fa in frame_analyses]
            captions = [fa['caption'] for fa in frame_analyses]
            
            # Count scene types
            from collections import Counter
            scene_counts = Counter(scenes)
            
            # Build summary
            summary_parts = []
            
            # Overall description - use "sampled frames" to be clear
            num_frames = len(frame_analyses)
            summary_parts.append(f"Video analyzed using {num_frames} sampled frames.")
            
            # Scene distribution - improved logic
            if scene_counts:
                # Get unique scene types
                unique_scenes = len(scene_counts)
                
                if unique_scenes == 1:
                    # All frames are same scene type
                    scene_type = list(scene_counts.keys())[0]
                    summary_parts.append(f"Consistent {scene_type} content throughout.")
                else:
                    # Multiple scene types
                    top_scenes = scene_counts.most_common(2)
                    
                    # Calculate percentages for better description
                    primary_scene, primary_count = top_scenes[0]
                    primary_pct = (primary_count / num_frames) * 100
                    
                    if primary_pct >= 75:
                        summary_parts.append(f"Predominantly {primary_scene} scenes ({primary_count}/{num_frames} frames).")
                    elif len(top_scenes) > 1:
                        secondary_scene, secondary_count = top_scenes[1]
                        summary_parts.append(f"Mixed content: {primary_count} {primary_scene}, {secondary_count} {secondary_scene} scenes.")
                    else:
                        summary_parts.append(f"Primarily {primary_scene} content.")
            
            # Key moments - improved description
            if captions:
                # Use first caption as opening
                summary_parts.append(f"Opening: {captions[0]}")
                
                # Add mid-point if available
                if len(captions) > 2:
                    mid_point = len(captions) // 2
                    if captions[mid_point] != captions[0]:  # Only if different
                        summary_parts.append(f"Mid-video: {captions[mid_point]}")
            
            return " ".join(summary_parts)
        
        except Exception as e:
            return f"Summary generation error: {e}"
    
    def search_video_frames_semantic(self, video_path: str, query: str, threshold: float = 0.20) -> List[Dict]:
        """Search video frames using CLIP semantic similarity"""
        try:
            if not self.image_analyzer or not self.image_analyzer.models_loaded:
                print("Image analyzer not available for semantic search")
                return []
            
            # Extract more frames for better coverage
            keyframes = self.extract_keyframes(video_path, num_frames=20)
            
            cap = cv2.VideoCapture(video_path)
            fps = cap.get(cv2.CAP_PROP_FPS)
            cap.release()
            
            matching_frames = []
            query_lower = query.lower()
            
            for frame_idx, frame in keyframes:
                frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                pil_image = Image.fromarray(frame_rgb)
                
                # Generate caption for context
                caption = self.image_analyzer.generate_caption(pil_image)
                
                # Compute CLIP semantic similarity
                clip_score = self.image_analyzer.compute_clip_similarity(pil_image, query)
                
                # Also check traditional keyword matching for hybrid approach
                keyword_match = 0
                if query_lower in caption.lower():
                    keyword_match = 1
                
                # Verify content match
                is_verified, verification_conf = self.verify_query_match_in_caption(caption, query)
                
                # Combined score with verification
                if is_verified:
                    combined_score = (clip_score * 0.7) + (keyword_match * 0.15) + (verification_conf * 0.15)
                else:
                    combined_score = (clip_score * 0.5) + (keyword_match * 0.2)  # Penalize unverified
                
                if combined_score >= threshold and is_verified:  # Require verification
                    timestamp = frame_idx / fps if fps > 0 else 0
                    matching_frames.append({
                        'frame_index': int(frame_idx),
                        'timestamp': round(timestamp, 2),
                        'caption': caption,
                        'clip_score': round(clip_score, 3),
                        'match_score': round(combined_score, 3),
                        'frame': pil_image
                    })
            
            # Sort by combined match score
            matching_frames.sort(key=lambda x: x['match_score'], reverse=True)
            return matching_frames
        
        except Exception as e:
            print(f"Frame search error: {e}")
            return []
    def enhance_query_for_search(self, query: str) -> str:
            query_lower = query.lower()
            
            if 'person' in query_lower or 'people' in query_lower or 'human' in query_lower:
                # Make person queries more explicit
                if 'in' in query_lower or 'inside' in query_lower:
                    return query  # Already specific
                else:
                    return f"a photograph showing {query}"
            
            elif 'car' in query_lower or 'vehicle' in query_lower:
                return f"a photograph showing {query}"
            
            elif 'animal' in query_lower or 'dog' in query_lower or 'cat' in query_lower:
                return f"a photograph showing {query}"
            
            elif len(query.split()) <= 2:
                return f"a photograph showing {query}"
            
            return query

    def verify_query_match_in_caption(self, caption: str, query: str) -> Tuple[bool, float]:
            query_lower = query.lower()
            caption_lower = caption.lower()
            
            query_words = set(query_lower.split())
            
            # Define critical search terms that MUST appear
            critical_terms = {
                'person': ['person', 'people', 'man', 'woman', 'human', 'individual', 'group', 'crowd', 'boy', 'girl', 'child'],
                'snow': ['snow', 'snowy', 'snowfall', 'winter', 'ice', 'icy', 'frozen', 'blizzard', 'sleet'],
                'car': ['car', 'vehicle', 'automobile', 'sedan', 'suv', 'truck', 'driving', 'traffic'],
                'dog': ['dog', 'puppy', 'canine', 'pet'],
                'cat': ['cat', 'kitten', 'feline', 'pet'],
                'building': ['building', 'structure', 'architecture', 'tower', 'skyscraper', 'construction'],
                'tree': ['tree', 'forest', 'woods', 'foliage', 'vegetation', 'palm', 'pine'],
                'water': ['water', 'ocean', 'sea', 'lake', 'river', 'stream', 'pond', 'waves'],
                'mountain': ['mountain', 'hill', 'peak', 'summit', 'mountainous', 'alpine'],
                'sky': ['sky', 'cloud', 'clouds', 'sunset', 'sunrise', 'horizon', 'cloudy'],
                'food': ['food', 'meal', 'dish', 'plate', 'eating', 'restaurant', 'dining'],
                'indoor': ['indoor', 'inside', 'interior', 'room', 'office', 'kitchen', 'bedroom'],
                'outdoor': ['outdoor', 'outside', 'exterior', 'street', 'road', 'sidewalk', 'path'],
                'rain': ['rain', 'rainy', 'rainfall', 'wet', 'umbrella', 'storm', 'pouring'],
                'beach': ['beach', 'sand', 'shore', 'coast', 'seaside', 'sandy'],
                'city': ['city', 'urban', 'downtown', 'metropolitan', 'skyline', 'buildings'],
                'nature': ['nature', 'natural', 'wilderness', 'landscape', 'scenic'],
                'animal': ['animal', 'wildlife', 'creature', 'beast'],
                'sports': ['sports', 'playing', 'game', 'athletic', 'exercise', 'running', 'soccer', 'basketball'],
                'technology': ['technology', 'computer', 'phone', 'device', 'screen', 'electronic', 'digital'],
                'traffic': ['traffic', 'cars', 'vehicles', 'road', 'highway', 'intersection', 'stoplight', 'signal'],
            }
            
            # Check if query contains any critical terms
            query_critical_entities = []
            for key_term, synonyms in critical_terms.items():
                if any(word in query_words for word in synonyms):
                    query_critical_entities.append((key_term, synonyms))
            
            if query_critical_entities:
                match_count = 0
                total_checks = len(query_critical_entities)
                
                for key_term, synonyms in query_critical_entities:
                    # Check if ANY synonym appears in caption
                    if any(syn in caption_lower for syn in synonyms):
                        match_count += 1
                
                # Calculate verification confidence
                verification_confidence = match_count / total_checks
                
                # Require at least 50% of critical terms to match
                if verification_confidence >= 0.5:
                    return True, verification_confidence
                else:
                    return False, verification_confidence
            
            # If no critical terms, fall back to basic keyword matching
            # Extract meaningful words from query (remove stop words)
            stop_words = {'a', 'an', 'the', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 
                          'find', 'show', 'me', 'all', 'video', 'videos', 'contains', 'containing'}
            query_keywords = [w for w in query_words if w not in stop_words and len(w) > 2]
            
            if not query_keywords:
                return True, 1.0  # No specific requirements, accept
            
            # Check how many query keywords appear in caption
            matches = sum(1 for word in query_keywords if word in caption_lower)
            keyword_confidence = matches / len(query_keywords)
            
            # Require at least 30% keyword match for general queries
            if keyword_confidence >= 0.3:
                return True, keyword_confidence
            else:
                return False, keyword_confidence    
    
    def compute_video_semantic_score(self, video_path: str, query: str) -> Tuple[float, List[Dict]]:
            try:
                if not self.image_analyzer or not self.image_analyzer.models_loaded:
                    return 0.0, []
                
                # Extract keyframes for analysis
                keyframes = self.extract_keyframes(video_path, num_frames=10)
                
                if not keyframes:
                    return 0.0, []
                
                # Enhance query for better CLIP matching
                enhanced_query = self.enhance_query_for_search(query)
                
                frame_scores = []
                verified_matches = 0
                total_verification_confidence = 0
                
                for frame_idx, frame in keyframes:
                    frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                    pil_image = Image.fromarray(frame_rgb)
                    
                    # Stage 1: Get CLIP similarity score
                    clip_score = self.image_analyzer.compute_clip_similarity(pil_image, enhanced_query)
                    
                    # Stage 2: Generate caption and verify content match
                    caption = self.image_analyzer.generate_caption(pil_image)
                    is_verified, verification_conf = self.verify_query_match_in_caption(caption, query)
                    
                    # Combine CLIP score with verification
                    if is_verified:
                        verified_matches += 1
                        total_verification_confidence += verification_conf
                        # Boost score if content is verified
                        adjusted_score = clip_score * (0.7 + (0.3 * verification_conf))
                    else:
                        adjusted_score = clip_score * 0.3
                    
                    frame_scores.append({
                        'frame_idx': frame_idx,
                        'score': clip_score,
                        'adjusted_score': adjusted_score,
                        'verified': is_verified,
                        'verification_conf': verification_conf,
                        'caption': caption
                    })
                
                # Calculate metrics on adjusted scores
                adjusted_scores = [f['adjusted_score'] for f in frame_scores]
                avg_adjusted = sum(adjusted_scores) / len(adjusted_scores)
                max_adjusted = max(adjusted_scores)
                
                # Calculate verification ratio
                verification_ratio = verified_matches / len(frame_scores)
                avg_verification_conf = total_verification_confidence / max(verified_matches, 1)
                
                # Stage 3: Final scoring with strict verification requirements
                if verification_ratio >= 0.5:
                    # At least 50% of frames verified - good match
                    final_score = (max_adjusted * 0.4) + (avg_adjusted * 0.6)
                    # Boost based on verification confidence
                    final_score = final_score * (0.8 + (0.2 * avg_verification_conf))
                elif verification_ratio >= 0.3:
                    # 30-50% frames verified - moderate match
                    final_score = (max_adjusted * 0.3) + (avg_adjusted * 0.7)
                    final_score = final_score * 0.9  # Small penalty
                elif verification_ratio >= 0.2:
                    # 20-30% frames verified - weak match
                    final_score = avg_adjusted * 0.8
                else:
                    # Less than 20% verified - likely false positive
                    final_score = avg_adjusted * 0.5  # Heavy penalty
                
                # Additional penalty if NO frames were verified
                if verified_matches == 0:
                    final_score = final_score * 0.3
                
                # Debug output
                print(f"Video scoring: query='{query}', verified={verified_matches}/{len(frame_scores)} frames, "
                      f"verification_ratio={verification_ratio:.2f}, final_score={final_score:.3f}")
                
                return round(final_score, 3), frame_scores
            
            except Exception as e:
                print(f"Video semantic score error: {e}")
                return 0.0, []

class DocumentAnalyzer:
    """Enhanced document analysis with NER, classification, and summarization"""
    
    def __init__(self):
        self.models_loaded = False
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        
        if not DOCUMENT_PROCESSING_AVAILABLE:
            print("⚠️  Document processing not available")
            return
        
        try:
            if TRANSFORMERS_AVAILABLE:
                print("Loading document analysis models...")
                
                # Summarization
                self.summarizer = pipeline(
                    "summarization", 
                    model="facebook/bart-large-cnn", 
                    device=0 if self.device == "cuda" else -1
                )
                
                # Sentiment Analysis
                self.sentiment_analyzer = pipeline(
                    "sentiment-analysis",
                    device=0 if self.device == "cuda" else -1
                )
                
                # Named Entity Recognition (NER)
                self.ner_pipeline = pipeline(
                    "ner",
                    model="dslim/bert-base-NER",
                    aggregation_strategy="simple",
                    device=0 if self.device == "cuda" else -1
                )
                
                # Zero-shot Classification for document categories
                self.classifier = pipeline(
                    "zero-shot-classification",
                    model="facebook/bart-large-mnli",
                    device=0 if self.device == "cuda" else -1
                )
                
                self.models_loaded = True
                print("✅ Document models loaded (Summarization, Sentiment, NER, Classification)")
                
        except Exception as e:
            print(f"⚠️  Document model loading error: {e}")
            self.models_loaded = False
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF with better error handling"""
        try:
            text = ""
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as e:
                        print(f"Warning: Could not extract text from page {page_num}: {e}")
                        continue
            
            return text.strip()
        except Exception as e:
            return f"PDF extraction error: {e}"
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from DOCX with better formatting"""
        try:
            doc = docx.Document(docx_path)
            paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            return "\n".join(paragraphs).strip()
        except Exception as e:
            return f"DOCX extraction error: {e}"
    
    def extract_text_from_txt(self, txt_path: str) -> str:
        """Extract text from TXT with encoding detection"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(txt_path, 'r', encoding=encoding) as file:
                    return file.read().strip()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                return f"TXT extraction error: {e}"
        
        return "Could not decode text file with supported encodings"
    
    def extract_named_entities(self, text: str) -> List[Dict[str, Any]]:
        """Extract named entities using BERT-based NER"""
        if not self.models_loaded:
            return []
        
        try:
            # Limit text length for NER
            text_for_ner = text[:5000] if len(text) > 5000 else text
            
            entities = self.ner_pipeline(text_for_ner)
            
            # Group by entity type
            entity_groups = {}
            for entity in entities:
                entity_type = entity['entity_group']
                entity_text = entity['word']
                score = entity['score']
                
                if entity_type not in entity_groups:
                    entity_groups[entity_type] = []
                
                entity_groups[entity_type].append({
                    'text': entity_text,
                    'score': round(score, 3)
                })
            
            return entity_groups
        
        except Exception as e:
            print(f"NER error: {e}")
            return {}
    
    def classify_document(self, text: str) -> Dict[str, float]:
        """Classify document into categories using zero-shot classification"""
        if not self.models_loaded:
            return {}
        
        try:
            # Document categories
            candidate_labels = [
                "technical documentation",
                "business report",
                "legal contract",
                "marketing material",
                "research paper",
                "financial document",
                "product specification",
                "meeting notes",
                "email correspondence",
                "policy document"
            ]
            
            # Use first 1000 characters for classification
            text_for_classification = text[:1000] if len(text) > 1000 else text
            
            result = self.classifier(text_for_classification, candidate_labels, multi_label=False)
            
            # Return top 3 categories with scores
            classifications = {}
            for label, score in zip(result['labels'][:3], result['scores'][:3]):
                classifications[label] = round(score, 3)
            
            return classifications
        
        except Exception as e:
            print(f"Classification error: {e}")
            return {}
    
    def generate_summary(self, text: str, max_length: int = 130) -> str:
        """Generate summary with better length handling"""
        if not self.models_loaded:
            return "Summary generation not available"
        
        try:
            # BART works best with 1024 tokens or less
            words = text.split()
            
            if len(words) < 50:
                return "Document too short for summarization"
            
            # Use appropriate chunk size
            text_for_summary = ' '.join(words[:400])  # ~400 words
            
            summary_result = self.summarizer(
                text_for_summary,
                max_length=max_length,
                min_length=30,
                do_sample=False
            )
            
            return summary_result[0]['summary_text']
        
        except Exception as e:
            print(f"Summarization error: {e}")
            # Fallback: return first few sentences
            sentences = text.split('.')[:3]
            return '. '.join(sentences) + '.'
    
    def analyze_document(self, file_path: str, file_extension: str) -> Dict[str, Any]:
        """Comprehensive document analysis with NER and classification"""
        try:
            # Extract text based on file type
            if file_extension == '.pdf':
                text = self.extract_text_from_pdf(file_path)
            elif file_extension == '.docx':
                text = self.extract_text_from_docx(file_path)
            elif file_extension == '.txt':
                text = self.extract_text_from_txt(file_path)
            else:
                return {'error': 'Unsupported document type'}
            
            if not text or len(text) < 10:
                return {'error': 'No text extracted or text too short'}
            
            if text.startswith(('PDF extraction error', 'DOCX extraction error', 'TXT extraction error')):
                return {'error': text}
            
            # Basic statistics
            words = text.split()
            word_count = len(words)
            char_count = len(text)
            
            # Initialize results
            summary = "Processing..."
            sentiment = "neutral"
            sentiment_score = 0.0
            named_entities = {}
            document_categories = {}
            
            # AI analysis
            if self.models_loaded and len(text) > 50:
                try:
                    # Summarization
                    summary = self.generate_summary(text)
                    
                    # Sentiment (on summary for speed)
                    text_for_sentiment = summary if len(summary) < 512 else text[:512]
                    sentiment_result = self.sentiment_analyzer(text_for_sentiment)[0]
                    sentiment = sentiment_result['label'].lower()
                    sentiment_score = round(sentiment_result['score'], 3)
                    
                    # Named Entity Recognition
                    named_entities = self.extract_named_entities(text)
                    
                    # Document Classification
                    document_categories = self.classify_document(text)
                    
                except Exception as e:
                    print(f"Document AI analysis error: {e}")
                    summary = f"Analysis partially failed: {e}"
            
            return {
                'word_count': word_count,
                'char_count': char_count,
                'summary': summary,
                'sentiment': sentiment,
                'sentiment_score': sentiment_score,
                'text_preview': text[:500] + "..." if len(text) > 500 else text,
                'named_entities': named_entities,
                'document_categories': document_categories
            }
        
        except Exception as e:
            return {'error': str(e)}


# ============================================================================
# UI FUNCTIONS
# ============================================================================

def connect_to_infinia(endpoint, access_key, secret_key, bucket, region):
    """Connect to DDN Infinia"""
    global s3_manager
    
    if not s3_manager:
        s3_manager = InfiniaS3Manager()
    
    success, message = s3_manager.connect(endpoint, access_key, secret_key, bucket, region)
    
    if success:
        status_html = f"""
        <div style="padding: 24px; background: linear-gradient(135deg, #1e7e34 0%, #28a745 100%); 
                    border-radius: 12px; text-align: center; border: 2px solid #1e7e34; box-shadow: 0 4px 12px rgba(40, 167, 69, 0.3);">
            <h3 style="color: white; margin: 0 0 12px 0; font-size: 20px; font-weight: 700;">✅ Connected</h3>
            <p style="color: white; font-size: 16px; margin: 8px 0; font-weight: 700; letter-spacing: 0.02em;">
                {bucket}
            </p>
            <p style="font-size: 13px; color: rgba(255, 255, 255, 0.95); margin: 4px 0 0 0; font-weight: 600; font-family: 'SF Mono', Consolas, monospace;">
                {endpoint}
            </p>
        </div>
        """
    else:
        status_html = f"""
        <div style="padding: 20px; background: #ffebee; border-radius: 10px; 
                    text-align: center; border: 2px solid #ef5350;">
            <h3 style="color: #c62828; margin: 0;">❌ Connection Failed</h3>
            <p style="color: #d32f2f; font-size: 14px; margin: 10px 0 0 0;">
                {message}
            </p>
        </div>
        """
    
    return status_html


def connect_to_exa_filesystem(host, username, port, key_file, passphrase, remote_path):
    """Connect to Exa Filesystem via SSH/SFTP"""
    global exa_manager
    
    if not EXA_FILESYSTEM_AVAILABLE:
        status_html = """
        <div style="padding: 20px; background: #ffebee; border-radius: 10px; 
                    text-align: center; border: 2px solid #ef5350;">
            <h3 style="color: #c62828; margin: 0;">❌ Exa Filesystem Not Available</h3>
            <p style="color: #d32f2f; font-size: 14px; margin: 10px 0 0 0;">
                Install paramiko: pip install paramiko
            </p>
        </div>
        """
        return status_html
    
    if not exa_manager:
        exa_manager = ExaFilesystemManager()
    
    # Handle key file - can be uploaded file or path
    key_path = None
    key_content = None
    
    if key_file:
        # If it's a file upload, read the content
        if hasattr(key_file, 'name'):
            # It's an uploaded file
            try:
                with open(key_file.name, 'r') as f:
                    key_content = f.read()
            except:
                with open(key_file, 'r') as f:
                    key_content = f.read()
        else:
            # It's a path string
            key_path = key_file
    
    success, message = exa_manager.connect(
        host=host,
        username=username,
        port=int(port) if port else 22,
        key_path=key_path,
        key_content=key_content,
        passphrase=passphrase if passphrase else None,
        remote_path=remote_path
    )
    
    if success:
        status_html = f"""
        <div style="padding: 24px; background: linear-gradient(135deg, #1e7e34 0%, #28a745 100%); 
                    border-radius: 12px; text-align: center; border: 2px solid #1e7e34; box-shadow: 0 4px 12px rgba(40, 167, 69, 0.3);">
            <h3 style="color: white; margin: 0 0 12px 0; font-size: 20px; font-weight: 700;">✅ Connected to Exa Filesystem</h3>
            <p style="color: white; font-size: 16px; margin: 8px 0; font-weight: 700; letter-spacing: 0.02em;">
                {remote_path}
            </p>
            <p style="font-size: 13px; color: rgba(255, 255, 255, 0.95); margin: 4px 0 0 0; font-weight: 600; font-family: 'SF Mono', Consolas, monospace;">
                {username}@{host}:{port if port else 22}
            </p>
        </div>
        """
    else:
        status_html = f"""
        <div style="padding: 20px; background: #ffebee; border-radius: 10px; 
                    text-align: center; border: 2px solid #ef5350;">
            <h3 style="color: #c62828; margin: 0;">❌ Connection Failed</h3>
            <p style="color: #d32f2f; font-size: 14px; margin: 10px 0 0 0;">
                {message}
            </p>
        </div>
        """
    
    return status_html


def upload_image(image_file, custom_caption):
    """Upload image with AI-generated metadata"""
    global s3_manager, ai_analyzer
    
    if not s3_manager or s3_manager.connection_status != "Connected":
        return None, "❌ Not connected to Infinia", ""
    
    if not image_file:
        return None, "❌ No image selected", ""
    
    try:
        # Ensure analyzer is available
        if not ai_analyzer or not ai_analyzer.models_loaded:
            return None, "❌ AI models not loaded. Please restart the application.", ""
        
        # Load image
        image = Image.open(image_file)
        
        # Analyze image
        analysis = ai_analyzer.analyze_image(image)
        
        # Prepare metadata
        final_caption = custom_caption if custom_caption else analysis['caption']
        
        metadata = {
            'modality': 'image',
            'caption': sanitize_for_s3_metadata(final_caption),
            'ai_caption': sanitize_for_s3_metadata(analysis['caption']),
            'width': str(analysis['width']),
            'height': str(analysis['height']),
            'aspect_ratio': str(analysis['aspect_ratio']),
            'dominant_scene': sanitize_for_s3_metadata(analysis.get('dominant_scene', 'unknown')),
            'detected_objects': sanitize_for_s3_metadata(analysis.get('detected_objects', 'unknown')),
            'upload_timestamp': datetime.now().isoformat()
        }
        
        # Add top 3 scenes
        if 'scene_classification' in analysis:
            for idx, scene in enumerate(analysis['scene_classification'][:3]):
                scene_text = f"{scene['label']} ({scene['score']:.3f})"
                metadata[f'scene_{idx+1}'] = sanitize_for_s3_metadata(scene_text)
        
        # Generate unique object key
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = os.path.basename(image_file)
        object_key = f"images/{timestamp}_{filename}"
        
        # Upload
        success, message = s3_manager.upload_file(
            image_file,
            object_key,
            metadata,
            content_type='image/jpeg'
        )
        
        if success:
            # Create metadata display
            metadata_html = f"""
            <div style="background: #e8f5e9; padding: 15px; border-radius: 8px; border-left: 4px solid #388e3c;">
                <h4 style="margin: 0 0 10px 0; color: #1b5e20;">📊 Image Metadata</h4>
                <p style="margin: 5px 0;"><strong>Caption:</strong> {final_caption}</p>
                <p style="margin: 5px 0;"><strong>AI Caption:</strong> {analysis['caption']}</p>
                <p style="margin: 5px 0;"><strong>Resolution:</strong> {analysis['width']} x {analysis['height']}</p>
                <p style="margin: 5px 0;"><strong>Dominant Scene:</strong> {analysis.get('dominant_scene', 'unknown')}</p>
                <p style="margin: 5px 0;"><strong>Detected Objects:</strong> {analysis.get('detected_objects', 'unknown')}</p>
                <p style="margin: 5px 0;"><strong>Object Key:</strong> <code>{object_key}</code></p>
            </div>
            """
            
            return image, f"✅ {message}", metadata_html
        else:
            return image, f"❌ {message}", ""
    
    except Exception as e:
        return None, f"❌ Upload error: {str(e)}", ""


def upload_video(video_file, custom_summary, custom_tags):
    """Upload video with AI-generated metadata, custom summary, custom tags, and PRE-COMPUTED embeddings"""
    global s3_manager, video_analyzer, ai_analyzer
    
    if not s3_manager or s3_manager.connection_status != "Connected":
        return None, "❌ Not connected to Infinia", ""
    
    if not video_file:
        return None, "❌ No video selected", ""
    
    try:
        # Ensure analyzer is available
        if not video_analyzer:
            return None, "❌ Video analyzer not initialized", ""
        
        if not ai_analyzer or not ai_analyzer.models_loaded:
            return None, "❌ AI models not loaded. Cannot generate embeddings.", ""
        
        # Analyze video
        print("📊 Analyzing video metadata...")
        analysis = video_analyzer.analyze_video(video_file)
        
        # Process custom tags
        tags_list = []
        if custom_tags and custom_tags.strip():
            tags_list = [tag.strip().lower() for tag in custom_tags.split(',') if tag.strip()]
        
        # Extract keyframes and generate CLIP embeddings
        print("🎞️ Extracting keyframes and generating embeddings...")
        keyframes = video_analyzer.extract_keyframes(video_file, num_frames=10)
        
        frame_embeddings = []
        for idx, (frame_idx, frame) in enumerate(keyframes):
            # Convert frame to PIL Image
            frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            pil_image = Image.fromarray(frame_rgb)
            
            # Generate CLIP embedding
            embedding = ai_analyzer.generate_clip_embedding(pil_image)
            
            # Generate caption for this frame
            caption = ai_analyzer.generate_caption(pil_image)
            
            # Store frame data
            frame_embeddings.append({
                'frame_idx': int(frame_idx),
                'embedding': embedding.tolist(),
                'caption': caption,
                'timestamp': round(frame_idx / analysis.get('fps', 30), 2) if analysis.get('fps', 0) > 0 else 0
            })
            
            print(f"  ✓ Frame {idx+1}/{len(keyframes)}: embedding generated ({len(embedding)} dims)")
        
        print(f"✅ Generated embeddings for {len(frame_embeddings)} keyframes")
        
        # Determine which summary to use
        ai_summary = analysis.get('video_summary', 'No summary')
        final_summary = custom_summary.strip() if custom_summary and custom_summary.strip() else ai_summary
        
        # Prepare metadata with embedding reference
        metadata = {
            'modality': 'video',
            'summary': sanitize_for_s3_metadata(final_summary),  # Use custom or AI summary
            'ai_summary': sanitize_for_s3_metadata(ai_summary),  # Always preserve AI summary
            'custom_summary': sanitize_for_s3_metadata(custom_summary.strip() if custom_summary and custom_summary.strip() else ''),  # Store custom summary separately
            'custom_tags': sanitize_for_s3_metadata(', '.join(tags_list)),  # Space after comma for better searchability
            'duration': str(analysis.get('duration', 0)),
            'fps': str(analysis.get('fps', 0)),
            'frame_count': str(analysis.get('frame_count', 0)),
            'resolution': sanitize_for_s3_metadata(analysis.get('resolution', 'unknown')),
            'keyframes_analyzed': str(len(frame_embeddings)),
            'has_embeddings': 'true',
            'embedding_version': 'v1',
            'upload_timestamp': datetime.now().isoformat()
        }
        
        # Add dominant scenes
        if 'dominant_scenes' in analysis:
            for idx, scene_info in enumerate(analysis['dominant_scenes'][:3]):
                scene_text = f"{scene_info['scene']} ({scene_info['count']} frames)"
                metadata[f'dominant_scene_{idx+1}'] = sanitize_for_s3_metadata(scene_text)
        
        # Generate unique object keys
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = os.path.basename(video_file)
        object_key = f"videos/{timestamp}_{filename}"
        embeddings_key = f"embeddings/{timestamp}_{filename}.json"
        
        # Upload embeddings as separate JSON object
        print(f"📤 Uploading embeddings to {embeddings_key}...")
        try:
            embeddings_data = json.dumps(frame_embeddings, indent=2).encode('utf-8')
            
            s3_manager.client.put_object(
                Bucket=s3_manager.bucket,
                Key=embeddings_key,
                Body=embeddings_data,
                ContentType='application/json',
                Metadata={
                    'video_key': object_key,
                    'num_frames': str(len(frame_embeddings)),
                    'embedding_dim': str(len(frame_embeddings[0]['embedding'])) if frame_embeddings else '0',
                    'created_at': datetime.now().isoformat()
                }
            )
            
            metadata['embeddings_key'] = embeddings_key
            metadata['embeddings_size_bytes'] = str(len(embeddings_data))
            
            print(f"✅ Embeddings uploaded successfully ({len(embeddings_data)} bytes)")
        
        except Exception as e:
            print(f"⚠️ Warning: Failed to upload embeddings: {e}")
            metadata['has_embeddings'] = 'false'
            metadata['embedding_error'] = str(e)
        
        # Upload video file with metadata
        print(f"📤 Uploading video to {object_key}...")
        success, message = s3_manager.upload_file(
            video_file,
            object_key,
            metadata,
            content_type='video/mp4'
        )
        
        if success:
            tags_display = ', '.join(tags_list) if tags_list else 'No custom tags'
            has_custom_summary = custom_summary and custom_summary.strip()
            
            # Build summary display
            summary_html = ""
            if has_custom_summary:
                summary_html = f"""
                <p style="margin: 5px 0;"><strong>📝 Custom Summary:</strong> {custom_summary.strip()}</p>
                <p style="margin: 5px 0;"><strong>🤖 AI Summary:</strong> <span style="color: #666;">{ai_summary}</span></p>
                """
            else:
                summary_html = f"""
                <p style="margin: 5px 0;"><strong>🤖 AI Summary:</strong> {ai_summary}</p>
                """
            
            metadata_html = f"""
            <div style="background: #e3f2fd; padding: 15px; border-radius: 8px; border-left: 4px solid #1976d2;">
                <h4 style="margin: 0 0 10px 0; color: #0d47a1;">📊 Video Metadata</h4>
                {summary_html}
                <p style="margin: 5px 0;"><strong>Custom Tags:</strong> <span style="background: #1976d2; color: white; padding: 2px 8px; border-radius: 3px;">{tags_display}</span></p>
                
                <div style="background: #c8e6c9; padding: 10px; border-radius: 5px; margin: 10px 0; border-left: 3px solid #4caf50;">
                    <p style="margin: 0 0 5px 0; font-weight: bold; color: #2e7d32;">⚡ Fast Search Enabled</p>
                    <p style="margin: 0; font-size: 12px; color: #1b5e20;">
                        ✓ {len(frame_embeddings)} keyframes indexed with CLIP embeddings<br/>
                        ✓ Embedding size: {metadata.get('embeddings_size_bytes', 'N/A')} bytes<br/>
                        ✓ Search will be 20-60x faster (no video download required)
                    </p>
                </div>
                
                <p style="margin: 5px 0;"><strong>Duration:</strong> {analysis.get('duration', 0)} seconds</p>
                <p style="margin: 5px 0;"><strong>Resolution:</strong> {analysis.get('resolution', 'unknown')}</p>
                <p style="margin: 5px 0;"><strong>FPS:</strong> {analysis.get('fps', 0)}</p>
                <p style="margin: 5px 0;"><strong>Frames:</strong> {analysis.get('frame_count', 0)}</p>
                <p style="margin: 5px 0;"><strong>Object Key:</strong> <code>{object_key}</code></p>
                <p style="margin: 5px 0;"><strong>Embeddings Key:</strong> <code>{embeddings_key}</code></p>
            </div>
            """
            
            return video_file, f"✅ {message}", metadata_html
        else:
            return video_file, f"❌ {message}", ""
    
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"❌ Upload error details:\n{error_details}")
        return None, f"❌ Upload error: {str(e)}", ""


def upload_document(document_file, custom_caption):
    """Upload document with enhanced AI analysis (NER, Classification, Summarization)"""
    global s3_manager, document_analyzer
    
    if not s3_manager or s3_manager.connection_status != "Connected":
        return "❌ Not connected to Infinia", ""
    
    if not document_file:
        return "❌ No document selected", ""
    
    try:
        if not document_analyzer:
            return "❌ Document analyzer not initialized", ""
        
        # Get file extension
        file_extension = os.path.splitext(document_file)[1].lower()
        
        # Analyze document
        analysis = document_analyzer.analyze_document(document_file, file_extension)
        
        if 'error' in analysis:
            return f"❌ Analysis error: {analysis['error']}", ""
        
        # Prepare metadata with sanitization
        final_caption = custom_caption if custom_caption else analysis.get('summary', 'No summary')
        
        metadata = {
            'modality': 'document',
            'caption': sanitize_for_s3_metadata(final_caption),
            'ai_summary': sanitize_for_s3_metadata(analysis.get('summary', 'No summary')),
            'word_count': str(analysis.get('word_count', 0)),
            'char_count': str(analysis.get('char_count', 0)),
            'sentiment': sanitize_for_s3_metadata(analysis.get('sentiment', 'unknown')),
            'sentiment_score': str(analysis.get('sentiment_score', 0)),
            'upload_timestamp': datetime.now().isoformat()
        }
        
        # Add document categories
        if 'document_categories' in analysis and analysis['document_categories']:
            for idx, (category, score) in enumerate(list(analysis['document_categories'].items())[:3]):
                metadata[f'category_{idx+1}'] = sanitize_for_s3_metadata(f"{category} ({score})")
        
        # Add named entities (top 3 per type)
        if 'named_entities' in analysis and analysis['named_entities']:
            entity_summary = []
            for entity_type, entities in list(analysis['named_entities'].items())[:5]:
                # Get unique entity texts
                unique_entities = list(set([e['text'] for e in entities[:3]]))
                entity_summary.append(f"{entity_type}: {', '.join(unique_entities[:3])}")
            
            metadata['named_entities'] = sanitize_for_s3_metadata(' | '.join(entity_summary))
        
        # Generate unique object key
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = os.path.basename(document_file)
        object_key = f"documents/{timestamp}_{filename}"
        
        # Upload
        success, message = s3_manager.upload_file(
            document_file,
            object_key,
            metadata,
            content_type='application/octet-stream'
        )
        
        if success:
            # Create enhanced metadata display
            metadata_html = f"""
            <div style="background: #e8f5e9; padding: 15px; border-radius: 8px; border-left: 4px solid #388e3c;">
                <h4 style="margin: 0 0 10px 0; color: #1b5e20;">📊 Document Metadata</h4>
                <p style="margin: 5px 0;"><strong>Caption:</strong> {final_caption}</p>
                <p style="margin: 5px 0;"><strong>Word Count:</strong> {metadata['word_count']}</p>
                <p style="margin: 5px 0;"><strong>Sentiment:</strong> {metadata['sentiment']} ({metadata['sentiment_score']})</p>
            """
            
            # Add document categories
            if 'document_categories' in analysis and analysis['document_categories']:
                metadata_html += '<div style="background: #fff3e0; padding: 10px; border-radius: 5px; margin: 10px 0;">'
                metadata_html += '<p style="margin: 0 0 5px 0; font-weight: bold; color: #e65100;">📂 Document Categories:</p>'
                for category, score in analysis['document_categories'].items():
                    confidence = "High" if score > 0.7 else "Medium" if score > 0.4 else "Low"
                    metadata_html += f'<p style="margin: 3px 0; font-size: 13px;">• {category.title()}: {score:.1%} <span style="color: #666;">({confidence})</span></p>'
                metadata_html += '</div>'
            
            # Add named entities
            if 'named_entities' in analysis and analysis['named_entities']:
                metadata_html += '<div style="background: #e1f5fe; padding: 10px; border-radius: 5px; margin: 10px 0;">'
                metadata_html += '<p style="margin: 0 0 5px 0; font-weight: bold; color: #01579b;">🏷️ Named Entities Detected:</p>'
                
                for entity_type, entities in analysis['named_entities'].items():
                    # Get top 5 unique entities
                    unique_entities = []
                    seen = set()
                    for entity in entities:
                        if entity['text'] not in seen:
                            unique_entities.append(entity)
                            seen.add(entity['text'])
                        if len(unique_entities) >= 5:
                            break
                    
                    if unique_entities:
                        entity_type_display = entity_type.replace('_', ' ').title()
                        metadata_html += f'<p style="margin: 5px 0; font-size: 13px;"><strong>{entity_type_display}:</strong> '
                        
                        entity_tags = []
                        for entity in unique_entities:
                            confidence_color = "#4caf50" if entity['score'] > 0.9 else "#ff9800" if entity['score'] > 0.7 else "#9e9e9e"
                            entity_tags.append(f'<span style="background: {confidence_color}; color: white; padding: 2px 6px; border-radius: 3px; font-size: 11px; margin: 2px;">{entity["text"]} ({entity["score"]:.2f})</span>')
                        
                        metadata_html += ' '.join(entity_tags)
                        metadata_html += '</p>'
                
                metadata_html += '</div>'
            
            metadata_html += f'<p style="margin: 10px 0 5px 0;"><strong>Object Key:</strong> <code>{object_key}</code></p>'
            metadata_html += "</div>"
            
            return f"✅ {message}", metadata_html
        else:
            return f"❌ {message}", ""
    
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"Upload error details:\n{error_details}")
        return f"❌ Upload error: {str(e)}", ""


def search_datasets(query, modality_filter, semantic_threshold=0.25):
    """Enhanced search with CLIP semantic filtering for videos and document previews"""
    global s3_manager, video_analyzer
    
    if not s3_manager or s3_manager.connection_status != "Connected":
        return "❌ Not connected to Infinia", [], [], [], []
    
    if not query:
        return "⚠️ Please enter a search query", [], [], [], []
    
    try:
        results = s3_manager.search_objects(query, modality_filter)
        
        if not results:
            return f"No results found for query: '{query}'", [], [], [], []
        
        # Create results table data
        table_data = []
        video_results = []
        document_results = []
        
        # Separate videos and documents for special processing
        video_candidates = []
        
        for result in results:
            metadata = result['metadata']
            modality = metadata.get('modality', 'unknown')
            
            if modality == 'video':
                video_candidates.append(result)
            elif modality == 'document':
                document_results.append(result)
            else:
                # For images, add directly to table
                icon_map = {'image': '📸', 'video': '🎬', 'document': '📄'}
                icon = icon_map.get(modality, '📦')
                
                caption = metadata.get('caption', metadata.get('summary', metadata.get('ai_caption', 'No caption')))
                if len(caption) > 100:
                    caption = caption[:97] + "..."
                
                table_data.append([
                    True,  # Select checkbox
                    result['key'],
                    f"{icon} {modality}",
                    caption,
                    f"{result['size']:,} bytes",
                    str(result['match_score']),
                ])
        
        # Apply semantic filtering to videos if analyzer is available
        if video_candidates and video_analyzer and video_analyzer.image_analyzer and video_analyzer.image_analyzer.models_loaded:
            print(f"Applying semantic filtering to {len(video_candidates)} videos with threshold {semantic_threshold}...")
            
            for video_result in video_candidates:
                try:
                    # Get video data
                    video_data = s3_manager.get_object(video_result['key'])
                    if not video_data:
                        continue
                    
                    # Save temporarily
                    temp_video = tempfile.NamedTemporaryFile(delete=False, suffix='.mp4')
                    temp_video.write(video_data)
                    temp_video.close()
                    
                    # Compute semantic score
                    semantic_score, frame_scores = video_analyzer.compute_video_semantic_score(temp_video.name, query)
                    
                    # Clean up
                    os.unlink(temp_video.name)
                    
                    # Only include video if it meets semantic threshold
                    if semantic_score >= semantic_threshold:
                        metadata = video_result['metadata']
                        caption = metadata.get('summary', metadata.get('ai_summary', 'No summary'))
                        if len(caption) > 100:
                            caption = caption[:97] + "..."
                        
                        # Update match score to semantic score
                        video_result['match_score'] = semantic_score
                        video_result['semantic_score'] = semantic_score
                        
                        table_data.append([
                            True,
                            video_result['key'],
                            "🎬 video",
                            caption,
                            f"{video_result['size']:,} bytes",
                            f"{semantic_score:.3f}",
                        ])
                        
                        video_results.append(video_result)
                        
                        # Enhanced debug output
                        print(f"✓ Video {video_result['key']}: score {semantic_score:.3f} (threshold: {semantic_threshold})")
                        if frame_scores:
                            verified_count = sum(1 for f in frame_scores if f.get('verified', False))
                            print(f"  → Verified frames: {verified_count}/{len(frame_scores)}")
                    else:
                        print(f"✗ Video {video_result['key']}: score {semantic_score:.3f} < threshold {semantic_threshold} - filtered out")
                        if frame_scores:
                            verified_count = sum(1 for f in frame_scores if f.get('verified', False))
                            print(f"  → Verified frames: {verified_count}/{len(frame_scores)} (insufficient match)")
                
                except Exception as e:
                    print(f"Error processing video {video_result['key']}: {e}")
                    continue
        else:
            # Fallback: use keyword matching for videos if CLIP not available
            print("Semantic filtering not available, using keyword matching for videos")
            for video_result in video_candidates:
                metadata = video_result['metadata']
                caption = metadata.get('summary', metadata.get('ai_summary', 'No summary'))
                if len(caption) > 100:
                    caption = caption[:97] + "..."
                
                table_data.append([
                    True,
                    video_result['key'],
                    "🎬 video",
                    caption,
                    f"{video_result['size']:,} bytes",
                    str(video_result['match_score']),
                ])
                
                video_results.append(video_result)
        
        # Add documents to table
        for doc_result in document_results:
            metadata = doc_result['metadata']
            caption = metadata.get('summary', metadata.get('ai_summary', 'No summary'))
            if len(caption) > 100:
                caption = caption[:97] + "..."
            
            table_data.append([
                True,
                doc_result['key'],
                "📄 document",
                caption,
                f"{doc_result['size']:,} bytes",
                str(doc_result['match_score']),
            ])
        
        # Sort video results by semantic score
        video_results.sort(key=lambda x: x.get('semantic_score', x['match_score']), reverse=True)
        
        # Sort document results by match score
        document_results.sort(key=lambda x: x['match_score'], reverse=True)
        
        # Prepare gallery - images only
        gallery_images = []
        for result in results[:20]:
            if result['metadata'].get('modality') == 'image':
                try:
                    image_data = s3_manager.get_object(result['key'])
                    if image_data:
                        image = Image.open(io.BytesIO(image_data))
                        gallery_images.append((image, result['key']))
                except:
                    continue
        
        # Prepare video cards for multiple videos
        video_cards = []
        for video_result in video_results[:10]:  # Show up to 10 videos
            try:
                video_data = s3_manager.get_object(video_result['key'])
                if video_data:
                    # Save temporarily
                    temp_video = tempfile.NamedTemporaryFile(delete=False, suffix='.mp4')
                    temp_video.write(video_data)
                    temp_video.close()
                    
                    metadata = video_result['metadata']
                    semantic_score = video_result.get('semantic_score', video_result['match_score'])
                    
                    video_cards.append({
                        'path': temp_video.name,
                        'key': video_result['key'],
                        'summary': metadata.get('summary', 'No summary'),
                        'duration': metadata.get('duration', '0'),
                        'resolution': metadata.get('resolution', 'unknown'),
                        'match_score': f"{semantic_score:.3f}"
                    })
            except Exception as e:
                print(f"Error preparing video card {video_result['key']}: {e}")
                continue
        
        # Prepare document cards for multiple documents
        document_cards = []
        for doc_result in document_results[:10]:  # Show up to 10 documents
            try:
                metadata = doc_result['metadata']
                
                # Extract key metadata
                doc_info = {
                    'key': doc_result['key'],
                    'summary': metadata.get('ai_summary', metadata.get('summary', 'No summary available')),
                    'word_count': metadata.get('word_count', 'N/A'),
                    'sentiment': metadata.get('sentiment', 'unknown'),
                    'sentiment_score': metadata.get('sentiment_score', '0'),
                    'match_score': doc_result['match_score'],
                    'categories': [],
                    'entities': []
                }
                
                # Extract categories
                for i in range(1, 4):
                    cat_key = f'category_{i}'
                    if cat_key in metadata:
                        doc_info['categories'].append(metadata[cat_key])
                
                # Extract entities
                if 'named_entities' in metadata:
                    entities_str = metadata['named_entities']
                    # Parse entity string: "ORG: DDN, NVIDIA | PER: John Smith"
                    if entities_str and entities_str != 'N/A':
                        doc_info['entities'] = entities_str.split(' | ')[:3]  # Top 3 entity types
                
                document_cards.append(doc_info)
                
            except Exception as e:
                print(f"Error preparing document card {doc_result['key']}: {e}")
                continue
        
        # Create summary HTML
        video_stats = ""
        if video_results:
            video_stats = f"""
            <div style="background: #e1bee7; padding: 10px; border-radius: 5px; margin: 10px 0;">
                <p style="margin: 3px 0; color: #4a148c;"><strong>🎬 Videos Found:</strong> {len(video_results)}</p>
                <p style="margin: 3px 0; color: #6a1b9a; font-size: 12px;">Using CLIP semantic similarity (threshold: {semantic_threshold})</p>
                <p style="margin: 3px 0; color: #6a1b9a; font-size: 11px;">Higher scores = better match | Adjust threshold slider for precision</p>
            </div>
            """
        
        document_stats = ""
        if document_results:
            document_stats = f"""
            <div style="background: #fff3e0; padding: 10px; border-radius: 5px; margin: 10px 0;">
                <p style="margin: 3px 0; color: #e65100;"><strong>📄 Documents Found:</strong> {len(document_results)}</p>
                <p style="margin: 3px 0; color: #f57c00; font-size: 12px;">View summaries, entities, and categories below</p>
            </div>
            """
        
        total_results = len(table_data)
        filtered_count = len(video_candidates) - len(video_results) if video_candidates else 0
        
        summary_html = f"""
        <div style="background: #ffebee; padding: 15px; border-radius: 8px; border-left: 4px solid #d32f2f; margin-bottom: 15px;">
            <h4 style="margin: 0 0 10px 0; color: #c62828;">🔍 Search Results: {total_results} items found</h4>
            <p style="margin: 5px 0; color: #666;">Query: <strong>{query}</strong></p>
            {video_stats}
            {document_stats}
            {f'<p style="margin: 5px 0; color: #ff5722; font-size: 12px;">⚠️ {filtered_count} videos filtered out (below semantic threshold)</p>' if filtered_count > 0 else ''}
            <p style="margin: 5px 0; color: #666; font-size: 12px;">
                ✅ All items selected by default. Uncheck rows you don't want to process.
            </p>
        </div>
        """
        
        return summary_html, table_data, gallery_images, video_cards, document_cards
    
    except Exception as e:
        return f"❌ Search error: {str(e)}", [], [], [], []


def search_datasets_fast(query, modality_filter, semantic_threshold=0.25):
    """
    FAST search using pre-computed embeddings - NO VIDEO DOWNLOAD
    
    This function searches videos using pre-computed CLIP embeddings stored during upload,
    eliminating the need to download and process videos during search time.
    
    Performance: 20-60x faster than original search_datasets function
    """
    global s3_manager, video_analyzer, ai_analyzer
    
    if not s3_manager or s3_manager.connection_status != "Connected":
        return "❌ Not connected to Infinia", [], [], [], []
    
    if not query:
        return "⚠️ Please enter a search query", [], [], [], []
    
    try:
        # Step 1: Metadata-based filtering (FAST - existing)
        print(f"🔍 Searching for: '{query}'")
        results = s3_manager.search_objects(query, modality_filter)
        
        if not results:
            return f"No results found for query: '{query}'", [], [], [], []
        
        # Step 2: Generate query embedding ONCE (not per video)
        if ai_analyzer and ai_analyzer.models_loaded:
            print(f"🧠 Generating query embedding...")
            query_embedding = ai_analyzer.generate_clip_text_embedding(query)
            print(f"✓ Query embedding generated ({len(query_embedding)} dimensions)")
        else:
            query_embedding = None
            print("⚠️ AI models not loaded, falling back to keyword matching")
        
        # Step 3: Separate results by modality
        table_data = []
        video_results = []
        document_results = []
        video_candidates = []
        
        for result in results:
            metadata = result['metadata']
            modality = metadata.get('modality', 'unknown')
            
            if modality == 'video':
                video_candidates.append(result)
            elif modality == 'document':
                document_results.append(result)
            else:
                # Images - add directly to table
                icon_map = {'image': '📸', 'video': '🎬', 'document': '📄'}
                icon = icon_map.get(modality, '📦')
                
                caption = metadata.get('caption', metadata.get('summary', metadata.get('ai_caption', 'No caption')))
                if len(caption) > 100:
                    caption = caption[:97] + "..."
                
                table_data.append([
                    True,
                    result['key'],
                    f"{icon} {modality}",
                    caption,
                    f"{result['size']:,} bytes",
                    str(result['match_score']),
                ])
        
        # Step 4: FAST video scoring using pre-computed embeddings
        if video_candidates and query_embedding is not None:
            print(f"⚡ Fast semantic scoring for {len(video_candidates)} videos using pre-computed embeddings...")
            
            videos_processed = 0
            videos_with_embeddings = 0
            videos_without_embeddings = 0
            
            for video_result in video_candidates:
                try:
                    metadata = video_result['metadata']
                    
                    # Check if embeddings are available
                    has_embeddings = metadata.get('has_embeddings') == 'true'
                    embeddings_key = metadata.get('embeddings_key')
                    
                    if has_embeddings and embeddings_key:
                        videos_with_embeddings += 1
                        
                        # Load pre-computed embeddings (FAST - only ~50KB)
                        try:
                            embeddings_data = s3_manager.get_object(embeddings_key)
                            if not embeddings_data:
                                print(f"⚠️ Embeddings file not found: {embeddings_key}")
                                videos_without_embeddings += 1
                                continue
                            
                            frame_embeddings = json.loads(embeddings_data.decode('utf-8'))
                            
                            # Compute similarities (FAST - pure math, no inference)
                            similarities = []
                            verified_count = 0
                            total_frames = len(frame_embeddings)
                            
                            for frame_data in frame_embeddings:
                                # Get pre-computed embedding
                                frame_emb = np.array(frame_data['embedding'])
                                
                                # Compute cosine similarity (dot product of normalized vectors)
                                similarity = float(np.dot(query_embedding, frame_emb))
                                
                                # Verify query-caption match
                                caption = frame_data.get('caption', '')
                                is_verified, verification_conf = video_analyzer.verify_query_match_in_caption(caption, query)
                                
                                # Apply verification boost/penalty
                                if is_verified:
                                    verified_count += 1
                                    adjusted_similarity = similarity * (1.0 + 0.2 * verification_conf)
                                else:
                                    adjusted_similarity = similarity * 0.8
                                
                                similarities.append({
                                    'raw': similarity,
                                    'adjusted': adjusted_similarity,
                                    'verified': is_verified,
                                    'caption': caption
                                })
                            
                            # Aggregate scores
                            if similarities:
                                # Extract adjusted scores
                                adjusted_scores = [s['adjusted'] for s in similarities]
                                
                                # Calculate aggregate metrics
                                max_score = max(adjusted_scores)
                                avg_score = np.mean(adjusted_scores)
                                verification_ratio = verified_count / total_frames
                                
                                # Final semantic score with verification weighting
                                if verification_ratio >= 0.5:
                                    semantic_score = (max_score * 0.4 + avg_score * 0.6) * (0.8 + 0.2 * verification_ratio)
                                elif verification_ratio >= 0.3:
                                    semantic_score = (max_score * 0.3 + avg_score * 0.7) * 0.9
                                elif verification_ratio >= 0.2:
                                    semantic_score = avg_score * 0.8
                                else:
                                    semantic_score = avg_score * 0.5
                                
                                if verified_count == 0:
                                    semantic_score *= 0.3
                                
                                print(f"  ✓ {video_result['key']}: score={semantic_score:.3f}, verified={verified_count}/{total_frames}")
                            else:
                                semantic_score = 0.0
                            
                            # Filter by threshold
                            if semantic_score >= semantic_threshold:
                                caption = metadata.get('summary', metadata.get('ai_summary', 'No summary'))
                                if len(caption) > 100:
                                    caption = caption[:97] + "..."
                                
                                video_result['match_score'] = semantic_score
                                video_result['semantic_score'] = semantic_score
                                
                                table_data.append([
                                    True,
                                    video_result['key'],
                                    "🎬 video",
                                    caption,
                                    f"{video_result['size']:,} bytes",
                                    f"{semantic_score:.3f}",
                                ])
                                
                                video_results.append(video_result)
                                videos_processed += 1
                            else:
                                print(f"  ✗ {video_result['key']}: score {semantic_score:.3f} below threshold")
                        
                        except json.JSONDecodeError as e:
                            print(f"⚠️ Failed to parse embeddings: {e}")
                            videos_without_embeddings += 1
                            continue
                    
                    else:
                        videos_without_embeddings += 1
                        print(f"  ⚠️ No embeddings for {video_result['key']}")
                
                except Exception as e:
                    print(f"❌ Error processing {video_result['key']}: {e}")
                    continue
            
            print(f"\n📊 Summary: {videos_with_embeddings} with embeddings, {videos_processed} passed threshold")
        
        else:
            print("⚠️ Fallback to keyword matching")
            for video_result in video_candidates:
                metadata = video_result['metadata']
                caption = metadata.get('summary', metadata.get('ai_summary', 'No summary'))
                if len(caption) > 100:
                    caption = caption[:97] + "..."
                
                table_data.append([
                    True,
                    video_result['key'],
                    "🎬 video",
                    caption,
                    f"{video_result['size']:,} bytes",
                    str(video_result['match_score']),
                ])
                video_results.append(video_result)
        
        # Sort results
        video_results.sort(key=lambda x: x.get('semantic_score', x['match_score']), reverse=True)
        document_results.sort(key=lambda x: x['match_score'], reverse=True)
        
        # Prepare gallery
        gallery_images = []
        for result in results[:20]:
            if result['metadata'].get('modality') == 'image':
                try:
                    image_data = s3_manager.get_object(result['key'])
                    if image_data:
                        image = Image.open(io.BytesIO(image_data))
                        gallery_images.append((image, result['key']))
                except:
                    continue
        
        # Prepare video cards
        video_cards = []
        for video_result in video_results[:10]:
            try:
                video_data = s3_manager.get_object(video_result['key'])
                if video_data:
                    temp_video = tempfile.NamedTemporaryFile(delete=False, suffix='.mp4')
                    temp_video.write(video_data)
                    temp_video.close()
                    
                    metadata = video_result['metadata']
                    semantic_score = video_result.get('semantic_score', video_result['match_score'])
                    
                    video_cards.append({
                        'path': temp_video.name,
                        'key': video_result['key'],
                        'summary': metadata.get('summary', 'No summary'),
                        'duration': metadata.get('duration', '0'),
                        'resolution': metadata.get('resolution', 'unknown'),
                        'match_score': f"{semantic_score:.3f}"
                    })
            except Exception as e:
                print(f"Error preparing video card: {e}")
                continue
        
        # Prepare document cards
        document_cards = []
        for doc_result in document_results[:10]:
            try:
                metadata = doc_result['metadata']
                doc_info = {
                    'key': doc_result['key'],
                    'summary': metadata.get('ai_summary', metadata.get('summary', 'No summary')),
                    'word_count': metadata.get('word_count', 'N/A'),
                    'sentiment': metadata.get('sentiment', 'unknown'),
                    'sentiment_score': metadata.get('sentiment_score', '0'),
                    'match_score': doc_result['match_score'],
                    'categories': [],
                    'entities': []
                }
                
                for i in range(1, 4):
                    cat_key = f'category_{i}'
                    if cat_key in metadata:
                        doc_info['categories'].append(metadata[cat_key])
                
                if 'named_entities' in metadata:
                    entities_str = metadata['named_entities']
                    if entities_str and entities_str != 'N/A':
                        doc_info['entities'] = entities_str.split(' | ')[:3]
                
                document_cards.append(doc_info)
            except Exception as e:
                print(f"Error preparing document card: {e}")
                continue
        
        # Create summary HTML
        video_stats = ""
        if video_results:
            video_stats = f"""
            <div style="background: #e1bee7; padding: 10px; border-radius: 5px; margin: 10px 0;">
                <p style="margin: 3px 0; color: #4a148c;"><strong>🎬 Videos:</strong> {len(video_results)}</p>
                <p style="margin: 3px 0; color: #6a1b9a; font-size: 12px;">⚡ FAST pre-computed embeddings</p>
            </div>
            """
        
        document_stats = ""
        if document_results:
            document_stats = f"""
            <div style="background: #fff3e0; padding: 10px; border-radius: 5px; margin: 10px 0;">
                <p style="margin: 3px 0; color: #e65100;"><strong>📄 Documents:</strong> {len(document_results)}</p>
            </div>
            """
        
        summary_html = f"""
        <div style="background: #ffebee; padding: 15px; border-radius: 8px; border-left: 4px solid #d32f2f;">
            <h4 style="margin: 0 0 10px 0; color: #c62828;">🔍 Search Results: {len(table_data)} items</h4>
            <p style="margin: 5px 0; color: #666;">Query: <strong>{query}</strong></p>
            {video_stats}
            {document_stats}
        </div>
        """
        
        return summary_html, table_data, gallery_images, video_cards, document_cards
    
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"❌ Search error:\n{error_details}")
        return f"❌ Search error: {str(e)}", [], [], [], []


def search_by_tags_only(query, modality_filter):
    """
    Search ONLY by custom tags - exact tag matching with no AI scoring
    Fast and precise for tag-based searches
    """
    global s3_manager
    
    if not s3_manager or s3_manager.connection_status != "Connected":
        return "❌ Not connected to Infinia", [], [], [], []
    
    if not query:
        return "⚠️ Please enter a search query", [], [], [], []
    
    try:
        print(f"🏷️ Searching by TAGS ONLY: '{query}'")
        
        # Extract search terms (remove stop words)
        stop_words = {'show', 'me', 'the', 'find', 'all', 'with', 'video', 'videos', 'image', 'images'}
        query_words = [word.lower().strip() for word in query.split() if word.lower() not in stop_words and len(word) > 2]
        
        if not query_words:
            query_words = [query.lower().strip()]
        
        print(f"🔍 Tag keywords: {query_words}")
        
        results = []
        response = s3_manager.client.list_objects_v2(Bucket=s3_manager.bucket)
        
        if 'Contents' not in response:
            return "No objects found in bucket", [], [], [], []
        
        # Search only in custom_tags field
        for obj in response['Contents']:
            try:
                head_response = s3_manager.client.head_object(
                    Bucket=s3_manager.bucket,
                    Key=obj['Key']
                )
                
                metadata = head_response.get('Metadata', {})
                
                # Filter by modality
                if modality_filter != "all":
                    obj_modality = metadata.get('modality', 'unknown')
                    if obj_modality != modality_filter:
                        continue
                
                # Check custom_tags field only
                custom_tags = metadata.get('custom_tags', '')
                if not custom_tags:
                    continue
                
                # Split tags and match
                individual_tags = [t.strip().lower() for t in custom_tags.split(',')]
                match_count = 0
                matched_tags = []
                
                for keyword in query_words:
                    for tag in individual_tags:
                        if keyword == tag or keyword in tag:
                            match_count += 1
                            matched_tags.append(tag)
                            break
                
                if match_count > 0:
                    results.append({
                        'key': obj['Key'],
                        'size': obj['Size'],
                        'metadata': metadata,
                        'match_score': match_count,
                        'matched_tags': matched_tags
                    })
                    print(f"  ✓ {obj['Key']}: matched tags {matched_tags} (score: {match_count})")
            
            except Exception as e:
                continue
        
        # Sort by match score
        results.sort(key=lambda x: x['match_score'], reverse=True)
        
        print(f"📊 Found {len(results)} objects with matching tags")
        
        # Build output
        table_data = []
        video_results = []
        document_results = []
        gallery_images = []
        
        for result in results:
            metadata = result['metadata']
            modality = metadata.get('modality', 'unknown')
            
            icon_map = {'image': '📸', 'video': '🎬', 'document': '📄'}
            icon = icon_map.get(modality, '📦')
            
            caption = metadata.get('caption', metadata.get('summary', metadata.get('ai_summary', 'No caption')))
            if len(caption) > 100:
                caption = caption[:97] + "..."
            
            table_data.append([
                True,
                result['key'],
                f"{icon} {modality}",
                caption,
                f"{result['size']:,} bytes",
                f"{result['match_score']} tags",
            ])
            
            if modality == 'video':
                video_results.append(result)
            elif modality == 'document':
                document_results.append(result)
            elif modality == 'image':
                try:
                    image_data = s3_manager.get_object(result['key'])
                    if image_data:
                        image = Image.open(io.BytesIO(image_data))
                        gallery_images.append((image, result['key']))
                except:
                    pass
        
        # Prepare video cards
        video_cards = []
        for video_result in video_results[:10]:
            try:
                video_data = s3_manager.get_object(video_result['key'])
                if video_data:
                    temp_video = tempfile.NamedTemporaryFile(delete=False, suffix='.mp4')
                    temp_video.write(video_data)
                    temp_video.close()
                    
                    metadata = video_result['metadata']
                    
                    video_cards.append({
                        'path': temp_video.name,
                        'key': video_result['key'],
                        'summary': metadata.get('summary', 'No summary'),
                        'duration': metadata.get('duration', '0'),
                        'resolution': metadata.get('resolution', 'unknown'),
                        'match_score': f"{video_result['match_score']} tags: {', '.join(video_result.get('matched_tags', []))}"
                    })
            except Exception as e:
                print(f"Error preparing video card: {e}")
                continue
        
        # Prepare document cards
        document_cards = []
        for doc_result in document_results[:10]:
            try:
                metadata = doc_result['metadata']
                doc_info = {
                    'key': doc_result['key'],
                    'summary': metadata.get('ai_summary', metadata.get('summary', 'No summary')),
                    'word_count': metadata.get('word_count', 'N/A'),
                    'sentiment': metadata.get('sentiment', 'unknown'),
                    'sentiment_score': metadata.get('sentiment_score', '0'),
                    'match_score': doc_result['match_score'],
                    'categories': [],
                    'entities': []
                }
                
                for i in range(1, 4):
                    cat_key = f'category_{i}'
                    if cat_key in metadata:
                        doc_info['categories'].append(metadata[cat_key])
                
                if 'named_entities' in metadata:
                    entities_str = metadata['named_entities']
                    if entities_str and entities_str != 'N/A':
                        doc_info['entities'] = entities_str.split(' | ')[:3]
                
                document_cards.append(doc_info)
            except Exception as e:
                print(f"Error preparing document card: {e}")
                continue
        
        # Summary HTML
        summary_html = f"""
        <div style="background: #e8f5e9; padding: 15px; border-radius: 8px; border-left: 4px solid #4caf50;">
            <h4 style="margin: 0 0 10px 0; color: #2e7d32;">🏷️ Tag Search Results: {len(results)} items</h4>
            <p style="margin: 5px 0; color: #666;">Query: <strong>{query}</strong></p>
            <p style="margin: 5px 0; color: #1b5e20; font-size: 12px;">
                ✓ Searched custom tags only (exact tag matching)
            </p>
        </div>
        """
        
        return summary_html, table_data, gallery_images, video_cards, document_cards
    
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"❌ Tag search error:\n{error_details}")
        return f"❌ Tag search error: {str(e)}", [], [], [], []


def search_unified(query, modality_filter, semantic_threshold, search_mode):
    """
    Unified search function that routes to appropriate search method
    """
    if search_mode == "tags_only":
        return search_by_tags_only(query, modality_filter)
    elif search_mode == "combined":
        # First find by tags, then rank by semantic similarity
        print("🔄 Combined search: Tags + Semantic ranking")
        tag_results = search_by_tags_only(query, modality_filter)
        # If tags found, use semantic search on those results
        # For now, just use tags-only (you can enhance this later)
        return tag_results
    else:  # semantic (default)
        return search_datasets_fast(query, modality_filter, semantic_threshold)


def copy_selected_objects(search_table_data, destination_bucket):
    """Copy selected objects to destination bucket"""
    global s3_manager
    
    if not s3_manager or s3_manager.connection_status != "Connected":
        return "❌ Not connected to Infinia"
    
    if search_table_data is None or len(search_table_data) == 0:
        return "⚠️ No search results available. Please search first."
    
    if not destination_bucket or not destination_bucket.strip():
        return "⚠️ Please enter destination bucket name"
    
    destination_bucket = destination_bucket.strip()
    
    try:
        success_count = 0
        failed_count = 0
        skipped_count = 0
        errors = []
        
        for idx, row in enumerate(search_table_data):
            try:
                if not isinstance(row, (list, tuple)) or len(row) < 6:
                    errors.append(f"Row {idx}: Invalid structure")
                    failed_count += 1
                    continue
                
                is_selected = row[0]
                source_key = str(row[1])
                
                if not is_selected:
                    skipped_count += 1
                    continue
                
                if not source_key or len(source_key) < 5:
                    errors.append(f"Row {idx}: Invalid object key")
                    failed_count += 1
                    continue
                
                # Create destination key
                dest_key = destination_bucket + '/' + os.path.basename(source_key)
                
                # Copy object
                copy_source = {'Bucket': s3_manager.bucket, 'Key': source_key}
                
                head_response = s3_manager.client.head_object(
                    Bucket=s3_manager.bucket,
                    Key=source_key
                )
                original_metadata = head_response.get('Metadata', {})
                
                s3_manager.client.copy_object(
                    CopySource=copy_source,
                    Bucket=s3_manager.bucket,
                    Key=dest_key,
                    Metadata=original_metadata,
                    MetadataDirective='REPLACE'
                )
                success_count += 1
                
            except ClientError as e:
                failed_count += 1
                error_code = e.response.get('Error', {}).get('Code', 'Unknown')
                errors.append(f"{source_key}: {error_code}")
            except Exception as e:
                failed_count += 1
                errors.append(f"Row {idx}: {str(e)}")
        
        result_html = f"""
        <div style="background: #e8f5e9; padding: 15px; border-radius: 8px; border-left: 4px solid #388e3c;">
            <h4 style="margin: 0 0 10px 0; color: #1b5e20;">✅ Copy Operation Complete</h4>
            <p style="margin: 5px 0;"><strong>Successfully copied:</strong> {success_count} objects</p>
            <p style="margin: 5px 0;"><strong>Skipped (not selected):</strong> {skipped_count} objects</p>
            <p style="margin: 5px 0;"><strong>Failed:</strong> {failed_count} objects</p>
            <p style="margin: 5px 0;"><strong>Bucket:</strong> {s3_manager.bucket}</p>
            <p style="margin: 5px 0;"><strong>Destination Folder:</strong> {destination_bucket}</p>
        """
        
        if errors:
            result_html += "<details style='margin-top: 10px;'><summary style='cursor: pointer; color: #d32f2f;'>⚠️ View Errors</summary><ul style='margin: 10px 0;'>"
            for error in errors[:15]:
                result_html += f"<li style='font-size: 12px; color: #d32f2f; margin: 3px 0;'>{error}</li>"
            if len(errors) > 15:
                result_html += f"<li style='font-size: 12px; color: #999;'>... and {len(errors) - 15} more errors</li>"
            result_html += "</ul></details>"
        
        result_html += "</div>"
        return result_html
        
    except Exception as e:
        return f"❌ Copy operation failed: {str(e)}"


def move_selected_objects(search_table_data, destination_bucket):
    """Move selected objects to destination bucket"""
    global s3_manager
    
    if not s3_manager or s3_manager.connection_status != "Connected":
        return "❌ Not connected to Infinia"
    
    if search_table_data is None or len(search_table_data) == 0:
        return "⚠️ No search results available. Please search first."
    
    if not destination_bucket or not destination_bucket.strip():
        return "⚠️ Please enter destination bucket name"
    
    destination_bucket = destination_bucket.strip()
    
    try:
        success_count = 0
        failed_count = 0
        skipped_count = 0
        errors = []
        
        for idx, row in enumerate(search_table_data):
            try:
                if not isinstance(row, (list, tuple)) or len(row) < 6:
                    errors.append(f"Row {idx}: Invalid structure")
                    failed_count += 1
                    continue
                
                is_selected = row[0]
                source_key = str(row[1])
                
                if not is_selected:
                    skipped_count += 1
                    continue
                
                if not source_key or len(source_key) < 5:
                    errors.append(f"Row {idx}: Invalid object key")
                    failed_count += 1
                    continue
                
                dest_key = destination_bucket + '/' + os.path.basename(source_key)
                
                # Copy object
                copy_source = {'Bucket': s3_manager.bucket, 'Key': source_key}
                
                head_response = s3_manager.client.head_object(
                    Bucket=s3_manager.bucket,
                    Key=source_key
                )
                original_metadata = head_response.get('Metadata', {})
                
                s3_manager.client.copy_object(
                    CopySource=copy_source,
                    Bucket=s3_manager.bucket,
                    Key=dest_key,
                    Metadata=original_metadata,
                    MetadataDirective='REPLACE'
                )
                
                # Delete source
                s3_manager.client.delete_object(
                    Bucket=s3_manager.bucket,
                    Key=source_key
                )
                success_count += 1
                
            except ClientError as e:
                failed_count += 1
                error_code = e.response.get('Error', {}).get('Code', 'Unknown')
                errors.append(f"{source_key}: {error_code}")
            except Exception as e:
                failed_count += 1
                errors.append(f"Row {idx}: {str(e)}")
        
        result_html = f"""
        <div style="background: #fff3e0; padding: 15px; border-radius: 8px; border-left: 4px solid #f57c00;">
            <h4 style="margin: 0 0 10px 0; color: #e65100;">✅ Move Operation Complete</h4>
            <p style="margin: 5px 0;"><strong>Successfully moved:</strong> {success_count} objects</p>
            <p style="margin: 5px 0;"><strong>Skipped (not selected):</strong> {skipped_count} objects</p>
            <p style="margin: 5px 0;"><strong>Failed:</strong> {failed_count} objects</p>
            <p style="margin: 5px 0;"><strong>Bucket:</strong> {s3_manager.bucket}</p>
            <p style="margin: 5px 0;"><strong>Destination Folder:</strong> {destination_bucket}</p>
        """
        
        if errors:
            result_html += "<details style='margin-top: 10px;'><summary style='cursor: pointer; color: #d32f2f;'>⚠️ View Errors</summary><ul style='margin: 10px 0;'>"
            for error in errors[:15]:
                result_html += f"<li style='font-size: 12px; color: #d32f2f; margin: 3px 0;'>{error}</li>"
            if len(errors) > 15:
                result_html += f"<li style='font-size: 12px; color: #999;'>... and {len(errors) - 15} more errors</li>"
            result_html += "</ul></details>"
        
        result_html += "</div>"
        return result_html
        
    except Exception as e:
        return f"❌ Move operation failed: {str(e)}"


def download_selected_objects(search_table_data):
    """Download selected objects as ZIP"""
    global s3_manager
    
    if not s3_manager or s3_manager.connection_status != "Connected":
        return "❌ Not connected to Infinia", None
    
    if search_table_data is None or len(search_table_data) == 0:
        return "⚠️ No search results available. Please search first.", None
    
    try:
        import zipfile
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        zip_filename = f"/tmp/infinia_download_{timestamp}.zip"
        
        success_count = 0
        failed_count = 0
        skipped_count = 0
        
        with zipfile.ZipFile(zip_filename, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for idx, row in enumerate(search_table_data):
                try:
                    if not isinstance(row, (list, tuple)) or len(row) < 6:
                        failed_count += 1
                        continue
                    
                    is_selected = row[0]
                    object_key = str(row[1])
                    
                    if not is_selected:
                        skipped_count += 1
                        continue
                    
                    if not object_key or len(object_key) < 5:
                        failed_count += 1
                        continue
                    
                    obj_data = s3_manager.get_object(object_key)
                    if obj_data:
                        filename = os.path.basename(object_key)
                        zipf.writestr(filename, obj_data)
                        success_count += 1
                    else:
                        failed_count += 1
                        
                except Exception as e:
                    failed_count += 1
                    print(f"Failed to download row {idx}: {e}")
        
        status_html = f"""
        <div style="background: #e3f2fd; padding: 15px; border-radius: 8px; border-left: 4px solid #1976d2;">
            <h4 style="margin: 0 0 10px 0; color: #0d47a1;">📦 Download Package Created</h4>
            <p style="margin: 5px 0;"><strong>Files packaged:</strong> {success_count}</p>
            <p style="margin: 5px 0;"><strong>Skipped (not selected):</strong> {skipped_count}</p>
            <p style="margin: 5px 0;"><strong>Failed:</strong> {failed_count}</p>
            <p style="margin: 5px 0;"><em>Download will start automatically below</em></p>
        </div>
        """
        
        return status_html, zip_filename
        
    except Exception as e:
        return f"❌ Download failed: {str(e)}", None


def list_all_objects(modality_filter):
    """List all objects"""
    global s3_manager
    
    if not s3_manager or s3_manager.connection_status != "Connected":
        return "❌ Not connected to Infinia", gr.update(choices=[])
    
    try:
        objects = s3_manager.get_all_objects(modality_filter)
        
        if not objects:
            return "No objects found in bucket", gr.update(choices=[])
        
        html = f"""
        <div style="background: #fff3e0; padding: 15px; border-radius: 8px; border-left: 4px solid #e65100;">
            <h4 style="margin: 0 0 15px 0; color: #bf360c;">📁 All Objects ({len(objects)})</h4>
            <div style="max-height: 400px; overflow-y: auto;">
                <table style="width: 100%; border-collapse: collapse; background: white;">
                    <thead style="background: #ffccbc; position: sticky; top: 0;">
                        <tr>
                            <th style="padding: 10px; text-align: left; border: 1px solid #ffab91;">Type</th>
                            <th style="padding: 10px; text-align: left; border: 1px solid #ffab91;">Object Key</th>
                            <th style="padding: 10px; text-align: left; border: 1px solid #ffab91;">Size</th>
                            <th style="padding: 10px; text-align: left; border: 1px solid #ffab91;">Modified</th>
                        </tr>
                    </thead>
                    <tbody>
        """
        
        object_keys = []
        
        for obj in objects:
            modality = obj.get('modality', 'unknown')
            icon_map = {'image': '📸', 'video': '🎬', 'document': '📄'}
            icon = icon_map.get(modality, '📦')
            
            object_keys.append(obj['key'])
            
            html += f"""
                <tr style="border-bottom: 1px solid #ffccbc;">
                    <td style="padding: 8px; border: 1px solid #ffccbc;">{icon} {modality}</td>
                    <td style="padding: 8px; border: 1px solid #ffccbc; font-family: monospace; font-size: 12px;">
                        <strong>{obj['key']}</strong>
                    </td>
                    <td style="padding: 8px; border: 1px solid #ffccbc;">{obj['size']}</td>
                    <td style="padding: 8px; border: 1px solid #ffccbc; font-size: 12px;">
                        {obj['last_modified'].strftime('%Y-%m-%d %H:%M')}
                    </td>
                </tr>
            """
        
        html += """
                    </tbody>
                </table>
            </div>
        </div>
        """
        
        return html, gr.update(choices=object_keys, value=object_keys[0] if object_keys else None)
    
    except Exception as e:
        return f"❌ Error listing objects: {str(e)}", gr.update(choices=[])


def load_object_details(object_key):
    """Load and display object details"""
    global s3_manager
    
    if not s3_manager or s3_manager.connection_status != "Connected":
        return None, "❌ Not connected to Infinia", None
    
    if not object_key:
        return None, "⚠️ Please select an object key", None
    
    try:
        metadata = s3_manager.get_object_metadata(object_key)
        
        if not metadata:
            return None, "❌ Object not found", None
        
        modality = metadata.get('modality', 'unknown')
        
        metadata_html = f"""
        <div style="background: #e1f5fe; padding: 15px; border-radius: 8px; border-left: 4px solid #0277bd;">
            <h4 style="margin: 0 0 10px 0; color: #01579b;">📋 Object Metadata</h4>
            <p style="margin: 5px 0;"><strong>Object Key:</strong> <code>{object_key}</code></p>
            <p style="margin: 5px 0;"><strong>Modality:</strong> {modality}</p>
        """
        
        for key, value in metadata.items():
            if key != 'modality':
                metadata_html += f'<p style="margin: 5px 0;"><strong>{key}:</strong> {value}</p>'
        
        metadata_html += "</div>"
        
        object_data = s3_manager.get_object(object_key)
        
        if not object_data:
            return None, metadata_html, None
        
        if modality == 'image':
            image = Image.open(io.BytesIO(object_data))
            return image, metadata_html, None
        elif modality == 'video':
            temp_video = tempfile.NamedTemporaryFile(delete=False, suffix='.mp4')
            temp_video.write(object_data)
            temp_video.close()
            return None, metadata_html, temp_video.name
        else:
            return None, metadata_html, None
    
    except Exception as e:
        return None, f"❌ Error loading object: {str(e)}", None


def search_video_content_semantic(object_key, query, threshold):
    """Search video frames using CLIP semantic similarity"""
    global s3_manager, video_analyzer
    
    if not s3_manager or s3_manager.connection_status != "Connected":
        return "❌ Not connected to Infinia", []
    
    if not object_key or not query:
        return "⚠️ Please provide object key and search query", []
    
    try:
        video_data = s3_manager.get_object(object_key)
        
        if not video_data:
            return "❌ Video not found", []
        
        # Save video temporarily
        temp_video = tempfile.NamedTemporaryFile(delete=False, suffix='.mp4')
        temp_video.write(video_data)
        temp_video.close()
        
        # Initialize analyzer if needed
        if not video_analyzer:
            return "❌ Video analyzer not initialized", []
        
        # Search frames using CLIP semantic similarity
        matching_frames = video_analyzer.search_video_frames_semantic(temp_video.name, query, threshold)
        
        # Clean up
        os.unlink(temp_video.name)
        
        if not matching_frames:
            return f"No matching frames found for query: '{query}' (threshold: {threshold})", []
        
        # Create results HTML
        results_html = f"""
        <div style="background: #fce4ec; padding: 15px; border-radius: 8px; border-left: 4px solid #c2185b;">
            <h4 style="margin: 0 0 10px 0; color: #880e4f;">🎞️ Semantic Video Frame Search Results</h4>
            <p style="margin: 5px 0;"><strong>Query:</strong> {query}</p>
            <p style="margin: 5px 0;"><strong>Matching Frames:</strong> {len(matching_frames)}</p>
            <p style="margin: 5px 0; color: #666; font-size: 12px;"><em>Using CLIP semantic similarity with threshold {threshold}</em></p>
        """
        
        for i, frame_info in enumerate(matching_frames[:10]):
            results_html += f"""
            <div style="background: white; padding: 10px; margin: 10px 0; border-radius: 5px;">
                <p style="margin: 0;"><strong>Frame {frame_info['frame_index']}</strong> @ {frame_info['timestamp']}s</p>
                <p style="margin: 5px 0; font-size: 13px;">{frame_info['caption']}</p>
                <p style="margin: 5px 0; font-size: 12px; color: #c2185b;">
                    CLIP Score: {frame_info['clip_score']} | Combined Score: {frame_info['match_score']}
                </p>
            </div>
            """
        
        results_html += "</div>"
        
        # Prepare gallery
        gallery_images = [(frame_info['frame'], f"Frame {frame_info['frame_index']} @ {frame_info['timestamp']}s\n{frame_info['caption']}\nCLIP: {frame_info['clip_score']}") for frame_info in matching_frames]
        
        return results_html, gallery_images
    
    except Exception as e:
        return f"❌ Search error: {str(e)}", []


def preview_selected_video(object_key):
    """Load and display video preview with metadata when selected in Video Search tab"""
    global s3_manager
    
    if not s3_manager or s3_manager.connection_status != "Connected":
        return gr.update(visible=False, value=None), gr.update(visible=False, value="")
    
    if not object_key:
        return gr.update(visible=False, value=None), gr.update(visible=False, value="")
    
    try:
        # Get video metadata
        metadata = s3_manager.get_object_metadata(object_key)
        
        if not metadata:
            return gr.update(visible=False, value=None), gr.update(visible=False, value="")
        
        # Get video data
        video_data = s3_manager.get_object(object_key)
        
        if not video_data:
            return gr.update(visible=False, value=None), gr.update(visible=False, value="")
        
        # Save video temporarily for preview
        temp_video = tempfile.NamedTemporaryFile(delete=False, suffix='.mp4')
        temp_video.write(video_data)
        temp_video.close()
        
        # Extract metadata
        summary = metadata.get('summary', 'No summary available')
        duration = metadata.get('duration', 'N/A')
        resolution = metadata.get('resolution', 'N/A')
        frame_count = metadata.get('frame_count', 'N/A')
        fps = metadata.get('fps', 'N/A')
        
        # Create info HTML
        info_html = f'''
        <div style="background: linear-gradient(135deg, #f3e5f5 0%, #e1bee7 100%); 
                    padding: 20px; border-radius: 10px; border: 2px solid #9c27b0; height: 100%;">
            <h3 style="margin: 0 0 15px 0; color: #6a1b9a; font-size: 18px;">
                📹 Video Information
            </h3>
            
            <div style="background: white; padding: 12px; border-radius: 8px; margin-bottom: 12px;">
                <p style="margin: 0 0 8px 0; font-weight: bold; color: #7b1fa2;">Summary:</p>
                <p style="margin: 0; font-size: 13px; color: #333; line-height: 1.6;">{summary}</p>
            </div>
            
            <div style="background: white; padding: 12px; border-radius: 8px;">
                <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 10px;">
                    <div>
                        <p style="margin: 0 0 8px 0; font-size: 12px; color: #666;">
                            <strong style="color: #7b1fa2;">⏱️ Duration:</strong> {duration}s
                        </p>
                        <p style="margin: 0 0 8px 0; font-size: 12px; color: #666;">
                            <strong style="color: #7b1fa2;">📐 Resolution:</strong> {resolution}
                        </p>
                    </div>
                    <div>
                        <p style="margin: 0 0 8px 0; font-size: 12px; color: #666;">
                            <strong style="color: #7b1fa2;">🎞️ Frames:</strong> {frame_count}
                        </p>
                        <p style="margin: 0 0 8px 0; font-size: 12px; color: #666;">
                            <strong style="color: #7b1fa2;">🎬 FPS:</strong> {fps}
                        </p>
                    </div>
                </div>
            </div>
            
            <div style="margin-top: 12px; padding: 10px; background: #fff3e0; border-radius: 6px; border-left: 3px solid #ff6f00;">
                <p style="margin: 0; font-size: 11px; color: #e65100;">
                    💡 Use the search query below to find specific moments in this video
                </p>
            </div>
        </div>
        '''
        
        return gr.update(visible=True, value=temp_video.name), gr.update(visible=True, value=info_html)
    
    except Exception as e:
        error_html = f'''
        <div style="background: #ffebee; padding: 15px; border-radius: 8px; border-left: 4px solid #d32f2f;">
            <p style="margin: 0; color: #c62828;">❌ Error loading video preview: {str(e)}</p>
        </div>
        '''
        return gr.update(visible=False, value=None), gr.update(visible=True, value=error_html)

# ============================================================================
# UI CREATION
# ============================================================================

def create_ui():
    """Create enhanced Gradio interface with tabbed layout"""
    
    # Detect runtime device
    # Application mainly uses CUDA or CPU, ignoring MPS for now based on logs
    runtime_device = "GPU (CUDA)" if torch.cuda.is_available() else "CPU"
    
    # Apple-style Custom CSS with red theme
    custom_css = """
    /* ===== CSS Variables - Apple Design System ===== */
    :root {
        /* DDN Infinia Color Palette */
        --ddn-red: #D0021B;
        --ddn-black: #000000;
        --ddn-blue: #005493;  /* Deep Blue */
        --ddn-light-blue: #E3F2FD; /* Light Blue for messages/backgrounds */
        --ddn-light-blue-bg: #E3F2FD;
        --ddn-text-blue: #0277BD;
        
        /* Apple-style Variables */
        --apple-font: -apple-system, BlinkMacSystemFont, "SF Pro Display", "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
        --apple-bg: #FFFFFF;
        --apple-card-bg: #FFFFFF;
        --apple-text-primary: #1D1D1F;
        --apple-text-secondary: #86868B;
        --apple-border: rgba(0, 0, 0, 0.1);
        --apple-shadow-sm: 0 1px 2px rgba(0, 0, 0, 0.04);
        --apple-shadow-md: 0 4px 12px rgba(0, 0, 0, 0.08);
        --apple-radius: 12px;
        --apple-radius-lg: 20px;
        --apple-transition: all 0.2s cubic-bezier(0.25, 0.1, 0.25, 1);
    }
    
    body {
        font-family: var(--apple-font) !important;
        background-color: #F5F5F7 !important;
        color: var(--apple-text-primary) !important;
        -webkit-font-smoothing: antialiased;
    }
    
    /* ===== Main Container ===== */
    .gradio-container {
        font-family: var(--apple-font) !important;
        max-width: 1400px !important;
        margin: 0 auto !important;
        padding: 40px 20px !important;
    }
    
    /* ===== Typography ===== */
    h1, h2, h3, h4, h5, h6 {
        font-family: var(--apple-font) !important;
        letter-spacing: -0.01em !important;
        transform: translateY(-1px);
    }
    
    /* ===== Input Fields - Apple Style ===== */
    input, textarea, select {
        border-radius: var(--apple-radius) !important;
        border: 1px solid var(--apple-border) !important;
        padding: 12px 16px !important;
        font-size: 15px !important;
        background: var(--apple-card-bg) !important;
        color: var(--apple-text-primary) !important;
        box-shadow: var(--apple-shadow-sm) !important;
        transition: var(--apple-transition) !important;
    }
    
    input:focus, textarea:focus, select:focus {
        outline: none !important;
        border-color: var(--apple-red) !important;
        box-shadow: 0 0 0 4px rgba(211, 47, 47, 0.1), var(--apple-shadow-sm) !important;
    }
    
    /* ===== Cards & Containers ===== */
    .gr-box, .gr-form, .gr-panel {
        background: var(--apple-card-bg) !important;
        border-radius: var(--apple-radius-lg) !important;
        border: 1px solid var(--apple-border) !important;
        box-shadow: var(--apple-shadow-sm) !important;
        padding: 24px !important;
    }
    
    .gr-box:hover {
        box-shadow: var(--apple-shadow-md) !important;
    }
    
    /* ===== Tabs - Apple Segmented Control Style ===== */
    .tabs {
        border-bottom: 1px solid var(--apple-border) !important;
        margin-bottom: 32px !important;
        background: transparent !important;
        min-height: 800px !important;
    }
    
    /* Ensure all tab contents have consistent minimum height */
    .tab-content, .tabitem {
        min-height: 700px !important;
    }
    
    .tab-nav {
        border: none !important;
        gap: 8px !important;
    }
    
    .tab-nav button {
        border-radius: 8px !important;
        padding: 10px 20px !important;
        font-weight: 500 !important;
        color: var(--apple-text-secondary) !important;
        background: transparent !important;
        border: none !important;
        box-shadow: none !important;
    }
    
    .tab-nav button[aria-selected="true"] {
        background: var(--apple-red) !important;
        color: white !important;
        box-shadow: var(--apple-shadow-sm) !important;
    }
    
    .tab-nav button:hover:not([aria-selected="true"]) {
        background: rgba(0, 0, 0, 0.04) !important;
        color: var(--apple-text-primary) !important;
    }
    
    /* ===== Textbox/Textarea Styling ===== */
    textarea, .gr-textbox textarea, .gr-text-input textarea {
        scrollbar-width: none !important; /* Firefox */
        -ms-overflow-style: none !important; /* IE/Edge */
    }
    
    textarea::-webkit-scrollbar, .gr-textbox textarea::-webkit-scrollbar {
        display: none !important; /* Chrome/Safari */
    }
    
    /* Custom Tags Field - Apple-style tag appearance */
    #custom-tags-input textarea {
        font-family: -apple-system, BlinkMacSystemFont, "SF Pro Display", sans-serif !important;
        line-height: 2 !important;
        padding: 12px !important;
    }
    
    /* Gallery - Apple Grid ===== */
    .gr-gallery {
        gap: 16px !important;
    }
    
    .gr-gallery img {
        border-radius: var(--apple-radius) !important;
        box-shadow: var(--apple-shadow-sm) !important;
        transition: var(--apple-transition) !important;
    }
    
    .gr-gallery img:hover {
        transform: scale(1.02);
        box-shadow: var(--apple-shadow-md) !important;
    }
    
    /* ===== Video Player ===== */
    video {
        border-radius: var(--apple-radius) !important;
        box-shadow: var(--apple-shadow-md) !important;
    }
    
    /* ===== Dataframe/Table ===== */
    .gr-dataframe {
        border-radius: var(--apple-radius) !important;
        overflow: hidden !important;
        box-shadow: var(--apple-shadow-sm) !important;
        border: 1px solid var(--apple-border) !important;
    }
    
    .gr-dataframe th {
        background: var(--apple-bg) !important;
        font-weight: 600 !important;
        color: var(--apple-text-primary) !important;
        padding: 12px 16px !important;
    }
    
    .gr-dataframe td {
        padding: 12px 16px !important;
        border-bottom: 1px solid var(--apple-border) !important;
    }
    
    .gr-dataframe tr:hover {
        background: rgba(211, 47, 47, 0.04) !important;
    }
    
    /* ===== Slider - Apple Style ===== */
    input[type="range"] {
        -webkit-appearance: none !important;
        background: var(--apple-border) !important;
        height: 4px !important;
        border-radius: 2px !important;
    }
    
    input[type="range"]::-webkit-slider-thumb {
        -webkit-appearance: none !important;
        width: 20px !important;
        height: 20px !important;
        border-radius: 50% !important;
        background: var(--apple-red) !important;
        cursor: pointer !important;
        box-shadow: var(--apple-shadow-sm) !important;
    }
    
    input[type="range"]::-webkit-slider-thumb:hover {
        transform: scale(1.1);
        box-shadow: 0 0 0 8px rgba(211, 47, 47, 0.1) !important;
    }
    
    /* ===== Dropdown ===== */
    .gr-dropdown {
        border-radius: var(--apple-radius) !important;
    }
    
    /* ===== File Upload ===== */
    .gr-file-upload {
        border: 2px dashed var(--apple-border) !important;
        border-radius: var(--apple-radius-lg) !important;
        background: var(--apple-card-bg) !important;
        padding: 32px !important;
        transition: var(--apple-transition) !important;
    }
    
    .gr-file-upload:hover {
        border-color: var(--ddn-red) !important;
        background: rgba(208, 2, 27, 0.02) !important;
    }
    
    /* ===== Markdown Content ===== */
    .gr-markdown {
        color: var(--apple-text-secondary) !important;
        line-height: 1.6 !important;
    }
    
    .gr-markdown h1, .gr-markdown h2, .gr-markdown h3 {
        color: var(--apple-text-primary) !important;
        font-weight: 600 !important;
    }
    
    /* ===== Animations ===== */
    @keyframes fadeIn {
        from {
            opacity: 0;
            transform: translateY(10px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }
    
    .gradio-container > * {
        animation: fadeIn 0.4s ease-out;
    }
    
    /* ===== Scrollbar - Apple Style ===== */
    ::-webkit-scrollbar {
        width: 8px;
        height: 8px;
    }
    
    ::-webkit-scrollbar-track {
        background: transparent;
    }
    
    ::-webkit-scrollbar-thumb {
        background: rgba(0, 0, 0, 0.2);
        border-radius: 4px;
    }
    
    ::-webkit-scrollbar-thumb:hover {
        background: rgba(0, 0, 0, 0.3);
    }
    
    /* ===== Status Messages ===== */
    .gr-info {
        border-radius: var(--apple-radius) !important;
        padding: 16px 20px !important;
        background: rgba(208, 2, 27, 0.08) !important;
        border-left: 4px solid var(--ddn-red) !important;
    }
    
    /* ===== Hide Gradio Footer Buttons ===== */
    footer {
        display: none !important;
    }
    
    .footer {
        display: none !important;
    }
    
    /* Hide "Use via API", "Built with Gradio", "Settings" buttons */
    button[aria-label*="Use via API"],
    button[aria-label*="Built with Gradio"],
    button[aria-label*="Settings"],
    .gradio-button-api,
    .gradio-button-built,
    .gradio-button-settings {
        display: none !important;
    }
    """
    
    with gr.Blocks(
        css=custom_css, 
        title="Intelligent Dataset Retrieval Platform",
        theme=gr.themes.Base(
            primary_hue="red",
            secondary_hue="orange",
            neutral_hue="slate"
        )
    ) as demo:
        
        # Apple-style Header with Logo
        import base64
        from pathlib import Path
        
        # Load and encode logo
        logo_path = Path(__file__).parent / "logo.png"
        if logo_path.exists():
            with open(logo_path, "rb") as f:
                logo_base64 = base64.b64encode(f.read()).decode()
            logo_html = f'<img src="data:image/png;base64,{logo_base64}" alt="Logo" style="width: 56px; height: 56px; filter: drop-shadow(0 2px 8px rgba(0,0,0,0.2));" />'
        else:
            logo_html = ""
        
        gr.HTML(f"""
        <div style="background: linear-gradient(135deg, #2E2E2E 0%, #1A1A1A 100%); 
                    padding: 24px 40px; 
                    margin-bottom: 32px; 
                    border-radius: 12px;
                    border: 1px solid #D0021B;
                    box-shadow: 0 4px 16px rgba(0, 0, 0, 0.3);">
            <div style="display: flex; align-items: center; justify-content: space-between; max-width: 1400px; margin: 0 auto;">
                <div style="display: flex; align-items: center; gap: 20px;">
                    {logo_html}
                    <div>
                        <h1 style="color: white; font-size: 28px; font-weight: 700; margin: 0 0 6px 0; 
                                   letter-spacing: 0.02em; 
                                   text-transform: uppercase;
                                   font-family: -apple-system, BlinkMacSystemFont, 'SF Pro Display', sans-serif;">
                            ENTERPRISE MULTIMODAL SEMANTIC SEARCH
                        </h1>
                        <p style="margin: 0; font-size: 13px; color: rgba(255, 255, 255, 0.7);
                                  font-weight: 400; letter-spacing: 0.01em;">
                            Powered by <span style="color: #D0021B; font-weight: 600;">DDN INFINIA</span> & <span style="color: #76B900; font-weight: 600;">NVIDIA GPU Computing</span>
                        </p>
                    </div>
                </div>
                <div style="background: #D0021B; color: white; padding: 8px 16px; 
                           border-radius: 6px; font-size: 11px; font-weight: 700;
                           letter-spacing: 0.05em; white-space: nowrap; margin-left: 20px;">
                    V6.0 ENTERPRISE
                </div>
            </div>
        </div>
        """)
        
        # Main tabs
        with gr.Tabs():
            
            with gr.Tab("Workflow Architecture"):
                gr.HTML(f"""
                <style>
                    .workflow-container {{
                        font-family: -apple-system, BlinkMacSystemFont, "SF Pro Display", sans-serif;
                        max-width: 1000px;
                        margin: 0 auto;
                        padding: 40px 20px;
                        background: #ffffff;
                        border-radius: 20px;
                        box-shadow: 0 4px 24px rgba(0,0,0,0.04);
                    }}
                    
                    .workflow-row {{
                        display: flex;
                        justify-content: center;
                        align-items: center;
                        gap: 40px;
                        margin-bottom: 60px;
                        position: relative;
                    }}
                    
                    .workflow-card {{
                        background: linear-gradient(145deg, #ffffff, #f5f5f7);
                        border: 1px solid rgba(0,0,0,0.05);
                        border-radius: 16px;
                        padding: 24px;
                        width: 220px;
                        text-align: center;
                        box-shadow: 0 8px 16px rgba(0,0,0,0.06);
                        transition: all 0.3s cubic-bezier(0.25, 0.8, 0.25, 1);
                        position: relative;
                        z-index: 2;
                    }}
                    
                    .workflow-card:hover {{
                        transform: translateY(-5px);
                        box-shadow: 0 12px 24px rgba(211, 47, 47, 0.15);
                        border-color: rgba(211, 47, 47, 0.2);
                    }}
                    
                    .card-icon {{
                        font-size: 32px;
                        margin-bottom: 12px;
                        display: block;
                    }}
                    
                    .card-title {{
                        font-weight: 700;
                        font-size: 16px;
                        color: #1d1d1f;
                        margin-bottom: 6px;
                        letter-spacing: -0.01em;
                    }}
                    
                    .card-subtitle {{
                        font-size: 13px;
                        color: #86868b;
                        line-height: 1.4;
                    }}
                    
                    .connector-arrow {{
                        color: #d32f2f;
                        font-size: 24px;
                        opacity: 0.6;
                        font-weight: bold;
                    }}
                    
                    .connector-line-vertical {{
                        position: absolute;
                        height: 40px;
                        width: 2px;
                        background: #e0e0e0;
                        left: 50%;
                        bottom: -50px;
                    }}
                    
                    .connector-elbow-right {{
                        position: absolute;
                        top: 50%;
                        right: -40px;
                        width: 40px;
                        height: 120px;
                        border-top: 2px solid #e0e0e0;
                        border-right: 2px solid #e0e0e0;
                        border-top-right-radius: 12px;
                        z-index: 1;
                    }}
                    
                    .s3-connection {{
                        position: absolute;
                        top: 50%;
                        right: -30px;
                        width: 30px;
                        height: 156px; /* Connects Row 1 to Row 2 */
                        border-right: 2px dashed #d32f2f;
                        opacity: 0.4;
                    }}
                    
                    .s3-connection::after {{
                        content: '▼';
                        position: absolute;
                        bottom: -10px;
                        right: -6px;
                        font-size: 12px;
                        color: #d32f2f;
                    }}
                    
                    .models-legend {{
                        margin-top: 40px;
                        background: #fbfbfd;
                        border-radius: 16px;
                        padding: 32px;
                        border: 1px solid rgba(0,0,0,0.03);
                    }}
                    
                    .model-tag {{
                        display: inline-flex;
                        align-items: center;
                        background: white;
                        padding: 8px 16px;
                        border-radius: 20px;
                        margin: 6px;
                        box-shadow: 0 2px 8px rgba(0,0,0,0.04);
                        font-size: 13px;
                        color: #424245;
                        border: 1px solid rgba(0,0,0,0.05);
                    }}
                    
                    .model-name {{
                        font-weight: 700;
                        color: #d32f2f;
                        margin-right: 6px;
                    }}
                    
                    /* Custom Arrow logic for row 2 flow */
                    .flow-path-arrow {{
                        position: absolute;
                        right: 250px;
                        top: 50%;
                        font-size: 24px;
                        color: #d32f2f;
                        opacity: 0.6;
                    }}
                    
                    .specs-container {{
                        margin-top: 24px;
                        display: grid;
                        grid-template-columns: 1fr 1fr;
                        gap: 24px;
                        text-align: left;
                    }}
                    
                    .spec-box {{
                        background: white;
                        border-radius: 12px;
                        padding: 20px;
                        border: 1px solid rgba(0,0,0,0.05);
                    }}
                    
                    .spec-title {{
                        font-weight: 600;
                        color: #1d1d1f;
                        margin-bottom: 12px;
                        font-size: 14px;
                        display: flex;
                        align-items: center;
                    }}
                    
                    .spec-detail {{
                        font-size: 13px;
                        color: #424245;
                        margin-bottom: 8px;
                        line-height: 1.5;
                    }}
                    
                    .spec-label {{
                        color: #86868b;
                        font-weight: 500;
                        margin-right: 4px;
                    }}
                    
                    .code-snippet {{
                        font-family: 'SF Mono', Consolas, monospace;
                        background: #f5f5f7;
                        padding: 4px 8px;
                        border-radius: 4px;
                        font-size: 12px;
                        color: #d32f2f;
                    }}
                </style>
                
                <div class="workflow-container">
                    <div style="text-align: center; margin-bottom: 50px;">
                        <h2 style="font-size: 28px; color: #1d1d1f; margin-bottom: 10px;">System Architecture</h2>
                        <p style="color: #86868b; font-size: 17px;">End-to-end multimodal retrieval operational flow</p>
                    </div>

                    <!-- Row 1: Ingestion -->
                    <div class="workflow-row">
                        <div class="workflow-card">
                            <span class="card-icon"></span>
                            <div class="card-title">Upload Video</div>
                            <div class="card-subtitle">User Interface Video/Doc Input</div>
                        </div>
                        
                        <div class="connector-arrow">→</div>
                        
                        <div class="workflow-card">
                            <span class="card-icon"></span>
                            <div class="card-title">Pre-Compute Embeddings</div>
                            <div class="card-subtitle">Feature Extraction & Metadata Gen</div>
                        </div>
                        
                        <div class="connector-arrow">→</div>
                        
                        <div class="workflow-card" style="border-color: #d32f2f;">
                            <div style="display: flex; justify-content: center; width: 100%; margin-bottom: 12px;">
                                <img src="data:image/png;base64,{logo_base64}" style="width: 48px; height: 48px;" alt="DDN Logo" />
                            </div>
                            <div class="card-title">DDN INFINIA S3</div>
                            <div class="card-subtitle">High-Perf Object Storage</div>
                            <!-- Connection to next row -->
                            <div style="position: absolute; right: 50%; bottom: -60px; height: 60px; width: 2px; border-left: 2px dashed #d32f2f; opacity: 0.3;"></div>
                             <div style="position: absolute; right: 49%; bottom: -68px; color: #d32f2f; font-size: 12px; opacity: 0.5;">▼</div>
                        </div>
                    </div>
                    
                    <!-- Row 2: Retrieval -->
                    <div class="workflow-row">
                        <div class="workflow-card">
                            <span class="card-icon"></span>
                            <div class="card-title">Search Query</div>
                            <div class="card-subtitle">Natural Language Input</div>
                        </div>
                        
                        <div class="connector-arrow">→</div>
                        
                        <div class="workflow-card">
                            <span class="card-icon"></span>
                            <div class="card-title">Load Pre-Comp Embeddings</div>
                            <div class="card-subtitle">Fetch Index from Infinia via S3 Protocol</div>
                            <div class="connector-line-vertical"></div>
                             <div style="position: absolute; bottom: -68px; left: 49%; color: #d32f2f; font-size: 12px; opacity: 0.6;">▼</div>
                        </div>
                        
                         <!-- Spacer to align with S3 above -->
                        <div style="width: 220px; opacity: 0;"></div>
                    </div>
                    
                    <!-- Row 3: Matching -->
                    <div class="workflow-row">
                        <div style="width: 220px; opacity: 0;"></div> <!-- Spacer -->
                        
                        <div class="workflow-card" style="background: #fff0f0; border-color: rgba(211, 47, 47, 0.1);">
                            <span class="card-icon"></span>
                            <div class="card-title">Cosine Similarity + Verification</div>
                            <div class="card-subtitle">Rank & Filter Results</div>
                        </div>
                        
                        <div style="width: 220px; opacity: 0;"></div> <!-- Spacer -->
                    </div>

                    <!-- Models Legend -->
                    <div class="models-legend">
                        <h3 style="margin: 0 0 20px 0; font-size: 18px; color: #1d1d1f;">AI Models Stack <span style="font-size: 13px; background: #e3f2fd; color: #1565c0; padding: 2px 8px; border-radius: 10px; vertical-align: middle; margin-left: 10px;">Running on: {runtime_device}</span></h3>
                        <div style="display: flex; flex-wrap: wrap; justify-content: center;">
                            <div class="model-tag">
                                <span class="model-name">OpenCV (cv2)</span>
                                Frame extraction & analysis
                            </div>
                            <div class="model-tag">
                                <span class="model-name">CLIP</span>
                                Contrastive learning on 400M image-text pairs
                            </div>
                            <div class="model-tag">
                                <span class="model-name">BLIP</span>
                                Image Captioning
                            </div>
                            <div class="model-tag">
                                <span class="model-name">ViT</span>
                                Scene Classification
                            </div>
                        </div>
                        
                        <!-- Technical Specifications -->
                        <div class="specs-container">
                            <div class="spec-box">
                                <div class="spec-title">Embedding Generation</div>
                                <div class="spec-detail"><span class="spec-label">Model:</span> openai/clip-vit-base-patch32</div>
                                <div class="spec-detail"><span class="spec-label">Input:</span> PIL Image (224x224)</div>
                                <div class="spec-detail"><span class="spec-label">Output:</span> 512-dimensional normalized embedding vector</div>
                            </div>
                            
                            <div class="spec-box">
                                <div class="spec-title">
                                    <div style="margin-right: 8px; display: flex; align-items: center;"><img src="data:image/png;base64,{logo_base64}" style="width: 20px; height: 20px;" alt="DDN Logo" /></div>
                                    Storage Architecture
                                </div>
                                <div class="spec-detail"><span class="spec-label">Backend:</span> DDN INFINIA (S3-compatible)</div>
                                <div class="spec-detail" style="margin-top: 8px;"><span class="spec-label">Video File:</span> <span class="code-snippet">videos/{{timestamp}}_{{filename}}.mp4</span></div>
                                <div class="spec-detail" style="font-size: 12px; color: #666; margin-left: 10px;">Original content + HTTP header metadata (Sub-ms retrieval)</div>
                                
                                <div class="spec-detail" style="margin-top: 8px;"><span class="spec-label">Embeddings:</span> <span class="code-snippet">embeddings/{{timestamp}}_{{filename}}.json</span></div>
                                <div class="spec-detail" style="font-size: 12px; color: #666; margin-left: 10px;">Lightweight JSON (~145KB) array of frame objects</div>
                            </div>
                        </div>
                    </div>
                </div>
                """)
            
            
            # ===== CONNECTION TAB =====
            with gr.Tab("Connection"):
                with gr.Tabs():
                    # S3 Connection Tab
                    with gr.Tab("S3 (Infinia)"):
                        gr.HTML('<div class="section-header"><h2>DDN Infinia S3 Connection</h2></div>')
                        
                        with gr.Row():
                            with gr.Column(scale=2):
                                endpoint_input = gr.Textbox(
                                    label="Infinia Endpoint URL",
                                    value="https://ddn-ai-demo-env:8111",
                                    placeholder="https://your-infinia-endpoint.com"
                                )
                                
                                with gr.Row():
                                    access_key_input = gr.Textbox(
                                        label="Access Key ID",
                                        value="7JPYE54HPIKA51RWSPGZ",
                                        type="password"
                                    )
                                    secret_key_input = gr.Textbox(
                                        label="Secret Access Key",
                                        value="sV2xcuDocH2psx6xm2OXP85LWVJtnNkSmzRP0Dui",
                                        type="password"
                                    )
                                
                                with gr.Row():
                                    bucket_input = gr.Textbox(
                                        label="Bucket Name",
                                        value="intelligent-datasets",
                                        placeholder="my-bucket"
                                    )
                                    region_input = gr.Textbox(
                                        label="Region",
                                        value="us-east-1"
                                    )
                                
                                connect_btn = gr.Button("Connect to Infinia S3", variant="primary", size="lg")
                            
                            with gr.Column(scale=1):
                                connection_status = gr.HTML("""
                                <div style="padding: 32px 24px; background: #ffffff; border-radius: 16px; 
                                            text-align: center; border: 1px solid rgba(0, 0, 0, 0.08);
                                            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08);">
                                    <div style="width: 48px; height: 48px; margin: 0 auto 16px; 
                                                background: rgba(211, 47, 47, 0.1); border-radius: 50%;
                                                display: flex; align-items: center; justify-content: center;
                                                font-size: 24px;"></div>
                                    <h3 style="color: #1d1d1f; font-weight: 600; margin: 0 0 8px 0; 
                                               font-size: 18px; letter-spacing: -0.02em;">Not Connected</h3>
                                    <p style="color: #6e6e73; font-size: 14px; margin: 0; line-height: 1.5;">
                                        Enter credentials and click Connect
                                    </p>
                                </div>
                                """)
                    
                    # Exa Filesystem Connection Tab
                    with gr.Tab("Exa Filesystem"):
                        gr.HTML('<div class="section-header"><h2>EXAScaler Filesystem (SSH)</h2></div>')
                        
                        with gr.Row():
                            with gr.Column(scale=2):
                                exa_host_input = gr.Textbox(
                                    label="SSH Host/IP Address",
                                    value="10.36.63.136",
                                    placeholder="10.36.63.136 or hostname"
                                )
                                
                                with gr.Row():
                                    exa_username_input = gr.Textbox(
                                        label="SSH Username",
                                        value="root",
                                        placeholder="root"
                                    )
                                    exa_port_input = gr.Textbox(
                                        label="SSH Port",
                                        value="22",
                                        placeholder="22"
                                    )
                                
                                exa_key_file_input = gr.Textbox(
                                    label="SSH Private Key Path",
                                    value="~/.ssh/cluster1_key",
                                    placeholder="~/.ssh/cluster1_key or upload file below",
                                    info="Specify path or upload key file"
                                )
                                
                                exa_key_upload = gr.File(
                                    label="Or Upload SSH Key File (Optional)",
                                    file_types=[".pem", ".key", ""],
                                    type="filepath"
                                )
                                
                                exa_passphrase_input = gr.Textbox(
                                    label="SSH Key Passphrase",
                                    type="password",
                                    placeholder="Enter passphrase if key is encrypted"
                                )
                                
                                exa_remote_path_input = gr.Textbox(
                                    label="Remote Path",
                                    value="/quackfs/Infinia-Exa-Test",
                                    placeholder="/quackfs/Infinia-Exa-Test"
                                )
                                
                                exa_connect_btn = gr.Button("Connect to Exa Filesystem", variant="primary", size="lg")
                            
                            with gr.Column(scale=1):
                                exa_connection_status = gr.HTML("""
                                <div style="padding: 32px 24px; background: #ffffff; border-radius: 16px; 
                                            text-align: center; border: 1px solid rgba(0, 0, 0, 0.08);
                                            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08);">
                                    <div style="width: 48px; height: 48px; margin: 0 auto 16px; 
                                                background: rgba(211, 47, 47, 0.1); border-radius: 50%;
                                                display: flex; align-items: center; justify-content: center;
                                                font-size: 24px;"></div>
                                    <h3 style="color: #1d1d1f; font-weight: 600; margin: 0 0 8px 0; 
                                               font-size: 18px; letter-spacing: -0.02em;">Not Connected</h3>
                                    <p style="color: #6e6e73; font-size: 14px; margin: 0; line-height: 1.5;">
                                        Enter SSH credentials and click Connect
                                    </p>
                                </div>
                                """)
            
            
            # ===== UPLOAD TABS =====
            with gr.Tab("Upload"):
                with gr.Tabs():
                    # Image Upload
                    with gr.Tab("Image"):
                        gr.HTML('<div class="section-header"><h2>Upload Image</h2></div>')
                        
                        with gr.Row():
                            with gr.Column():
                                image_upload_file = gr.Image(label="Select Image", type="filepath")
                                image_custom_caption = gr.Textbox(label="Custom Caption (Optional)", placeholder="Leave empty for AI-generated caption")
                                upload_image_btn = gr.Button("Upload Image", variant="primary")
                            
                            with gr.Column():
                                upload_image_preview = gr.Image(label="Preview")
                                upload_image_status = gr.HTML()
                        
                        upload_image_metadata = gr.HTML()
                    
                    # Video Upload
                    with gr.Tab("Video"):
                        gr.HTML('<div class="section-header"><h2>Upload Video</h2></div>')
                        
                        with gr.Row():
                            with gr.Column():
                                video_upload_file = gr.Video(label="Select Video")
                                video_custom_summary = gr.Textbox(
                                    label="Custom Summary (Optional)", 
                                    placeholder="Enter your own summary or leave empty for AI-generated summary",
                                    lines=3
                                )
                                video_custom_tags = gr.Textbox(
                                    label="Custom Tags (Optional)", 
                                    placeholder="Enter comma-separated tags (e.g., outdoor, car, highway)",
                                    info="Tags help with quick filtering and search",
                                    elem_id="custom-tags-input"
                                )
                                upload_video_btn = gr.Button("Upload Video", variant="primary")
                            
                            with gr.Column():
                                upload_video_preview = gr.Video(label="Preview")
                                upload_video_status = gr.HTML()
                        
                        upload_video_metadata = gr.HTML()
                    
                    # Document Upload
                    with gr.Tab("Document"):
                        gr.HTML('<div class="section-header"><h2>Upload Document</h2></div>')
                        
                        with gr.Row():
                            with gr.Column():
                                document_upload_file = gr.File(label="Select Document (PDF, DOCX, TXT)")
                                document_custom_caption = gr.Textbox(label="Custom Caption (Optional)", placeholder="Leave empty for AI-generated summary")
                                upload_document_btn = gr.Button("Upload Document", variant="primary")
                            
                            with gr.Column():
                                upload_document_status = gr.HTML()
                                upload_document_metadata = gr.HTML()
            
            # ===== INTELLIGENT SEARCH TAB =====
            with gr.Tab("Intelligent Search"):
                gr.HTML('<div class="section-header"><h2>AI-Powered Dataset Search</h2></div>')
                
                gr.Markdown("""
                **Semantic Video Search with CLIP** - Videos are filtered using AI-powered semantic similarity matching.
                Adjust the threshold to control precision: higher = more accurate, lower = more results.
                """)
                
                with gr.Row():
                    search_query = gr.Textbox(
                        label="Search Query",
                        placeholder="e.g., 'find me all the videos that contain snow road'",
                        scale=3
                    )
                    search_mode = gr.Dropdown(
                        label="Search Mode",
                        choices=[
                            ("AI Search", "semantic"),
                            ("Tags Only", "tags_only"),
                            ("Combined", "combined")
                        ],
                        value="semantic",
                        scale=1
                    )
                    search_modality = gr.Dropdown(
                        label="Filter by Type",
                        choices=["all", "image", "video", "document"],
                        value="all",
                        scale=1
                    )
                    semantic_threshold = gr.Slider(
                        label="Video Semantic Match Threshold",
                        minimum=0.20,
                        maximum=0.70,
                        value=0.30,
                        step=0.05,
                        info="Recommended: 0.45-0.55 for specific objects/people, 0.30-0.40 for vehicles/scenes, 0.20-0.30 for broad categories"
                    )
                    search_btn = gr.Button("Search", variant="primary", scale=1)
                
                search_summary = gr.HTML()
                
                gr.Markdown("### Search Results Table")
                search_results_table = gr.Dataframe(
                    headers=["Select", "Object Key", "Type", "Caption/Summary", "Size", "Match Score"],
                    datatype=["bool", "str", "str", "str", "str", "str"],
                    col_count=(6, "fixed"),
                    interactive=True,
                    wrap=True
                )
                
                gr.Markdown("### Bulk Operations")
                with gr.Row():
                    destination_bucket = gr.Textbox(
                        label="Destination Folder/Bucket",
                        placeholder="e.g., 'my-backup-folder'",
                        scale=2
                    )
                    with gr.Column(scale=1):
                        copy_btn = gr.Button(" Copy Selected", variant="secondary")
                        move_btn = gr.Button(" Move Selected", variant="secondary")
                        download_btn = gr.Button(" Download Selected", variant="secondary")
                
                operation_status = gr.HTML()
                download_file = gr.File(label="Download", visible=False)
                
                gr.Markdown("### Video Results")
                gr.HTML("""
                <div style="background: #e3f2fd; padding: 12px; border-radius: 8px; margin: 10px 0; border: 1px solid #90caf9;">
                    <p style="margin: 0; color: #0d47a1; font-size: 14px; display: flex; align-items: center;">
                        <span style="font-weight: 600; margin-right: 6px;">Multiple Video Viewer:</span> All matching videos are displayed below with playback controls.
                    </p>
                </div>
                """)
                
                gr.Markdown("### Image Gallery")
                search_gallery = gr.Gallery(
                    label="Matching Images",
                    columns=4,
                    height="auto",
                    preview=True
                )
                
                # Multiple video players with metadata (up to 10) - Grid Layout
                video_info_list = []
                video_player_list = []
                
                for row_idx in range(5):
                    with gr.Row():
                        for col_idx in range(2):
                            with gr.Column():
                                with gr.Group():
                                    v_info = gr.HTML(visible=False)
                                    v_player = gr.Video(label="", interactive=False, visible=False)
                                    video_info_list.append(v_info)
                                    video_player_list.append(v_player)
                
                # Unpack for compatibility with event handlers
                video_info_1, video_info_2, video_info_3, video_info_4, video_info_5, \
                video_info_6, video_info_7, video_info_8, video_info_9, video_info_10 = video_info_list
                
                video_player_1, video_player_2, video_player_3, video_player_4, video_player_5, \
                video_player_6, video_player_7, video_player_8, video_player_9, video_player_10 = video_player_list
                
                gr.Markdown("### Document Results")
                gr.HTML("""
                <div style="background: #fff3e0; padding: 12px; border-radius: 8px; margin: 10px 0; border: 1px solid #ffe0b2;">
                    <p style="margin: 0; color: #e65100; font-size: 14px; display: flex; align-items: center;">
                        <span style="font-weight: 600; margin-right: 6px;">Document Previews:</span> View summaries, entities, and categories for all matching documents.
                    </p>
                </div>
                """)
                
                # Multiple document preview cards (up to 10)
                with gr.Row():
                    with gr.Column():
                        doc_card_1 = gr.HTML(visible=False)
                        doc_card_2 = gr.HTML(visible=False)
                        doc_card_3 = gr.HTML(visible=False)
                        doc_card_4 = gr.HTML(visible=False)
                        doc_card_5 = gr.HTML(visible=False)
                    
                    with gr.Column():
                        doc_card_6 = gr.HTML(visible=False)
                        doc_card_7 = gr.HTML(visible=False)
                        doc_card_8 = gr.HTML(visible=False)
                        doc_card_9 = gr.HTML(visible=False)
                        doc_card_10 = gr.HTML(visible=False)
            
            # ===== BROWSE DATASETS TAB =====
            with gr.Tab("Browse Datasets"):
                gr.HTML('<div class="section-header"><h2>Dataset Browser</h2></div>')
                
                with gr.Row():
                    browse_modality = gr.Dropdown(
                        label="Filter by Type",
                        choices=["all", "image", "video", "document"],
                        value="all"
                    )
                    list_objects_btn = gr.Button("List All Objects", variant="secondary")
                
                dataset_list_html = gr.HTML()
                
                gr.HTML('<div style="height: 20px;"></div>')
                gr.Markdown("### Load Object Details")
                
                with gr.Row():
                    with gr.Column(scale=2):
                        object_key_dropdown = gr.Dropdown(
                            label="Select Object Key",
                            choices=[],
                            interactive=True,
                            allow_custom_value=True
                        )
                    with gr.Column(scale=1):
                        load_object_btn = gr.Button("Load Details", variant="primary")
                
                with gr.Row():
                    with gr.Column():
                        object_image_preview = gr.Image(label="Image Preview")
                        object_metadata_display = gr.HTML()
                    with gr.Column():
                        object_video_preview = gr.Video(label="Video Preview")
            
            # ===== VIDEO SEARCH TAB =====
            with gr.Tab("Video Search & Summarization"):
                gr.HTML('<div class="section-header"><h2>Semantic Video Frame Search with NVIDIA CLIP</h2></div>')
                
                gr.Markdown("""
                **Advanced Video Search powered by NVIDIA CLIP**
                
                This feature uses CLIP (Contrastive Language-Image Pre-training) to perform semantic similarity matching between your text query and video frames.
                Unlike keyword matching, CLIP understands the meaning and context of your query to find visually relevant frames.
                
                **How it works:**
                1. Select a video from your dataset
                2. Enter a natural language query describing what you're looking for
                3. Adjust the similarity threshold (lower = more results, higher = more precise)
                4. CLIP analyzes frames and returns matches with confidence scores
                
                **Example queries:**
                - "a person sitting at a desk"
                - "outdoor mountain landscape"
                - "cars driving on a highway"
                - "people playing basketball"
                - "sunset over the ocean"
                """)
                
                with gr.Row():
                    video_search_key = gr.Dropdown(
                        label="Select Video",
                        choices=[],
                        interactive=True,
                        allow_custom_value=True,
                        scale=2
                    )
                    video_search_query = gr.Textbox(
                        label="Semantic Search Query",
                        placeholder="Describe what you want to find in the video...",
                        scale=2
                    )
                    video_search_threshold = gr.Slider(
                        label="Similarity Threshold",
                        minimum=0.1,
                        maximum=0.5,
                        value=0.20,
                        step=0.05,
                        scale=1,
                        info="Lower = more results"
                    )
                    video_search_btn = gr.Button("Search Frames", variant="primary", scale=1)
                # Video preview section (shown when video is selected)
                with gr.Row():
                    with gr.Column(scale=1):
                        video_preview_player = gr.Video(
                            label="Selected Video Preview",
                            visible=False
                        )
                    with gr.Column(scale=1):
                        video_preview_info = gr.HTML(visible=False)
                
                video_search_results = gr.HTML()
                video_frame_gallery = gr.Gallery(
                    label="Matching Frames with CLIP Scores",
                    columns=3,
                    height="auto"
                )
        

        # ===== EVENT HANDLERS =====
        
        # Connection
        connect_btn.click(
            fn=connect_to_infinia,
            inputs=[endpoint_input, access_key_input, secret_key_input, bucket_input, region_input],
            outputs=[connection_status]
        )
        
        # Image upload
        upload_image_btn.click(
            fn=upload_image,
            inputs=[image_upload_file, image_custom_caption],
            outputs=[upload_image_preview, upload_image_status, upload_image_metadata]
        )
        
        # Video upload
        upload_video_btn.click(
            fn=upload_video,
            inputs=[video_upload_file, video_custom_summary, video_custom_tags],
            outputs=[upload_video_preview, upload_video_status, upload_video_metadata]
        )
        
        # Document upload
        upload_document_btn.click(
            fn=upload_document,
            inputs=[document_upload_file, document_custom_caption],
            outputs=[upload_document_status, upload_document_metadata]
        )
        
        # Search - now returns video_cards and document_cards
        def search_and_display_videos(query, modality, threshold, search_mode):
            summary, table, gallery, video_cards, document_cards = search_unified(query, modality, threshold, search_mode)
            
            # Prepare outputs: 10 video info HTML + 10 video players + 10 document cards
            video_info_outputs = [gr.update(visible=False)] * 10
            video_player_outputs = [gr.update(visible=False, value=None)] * 10
            document_card_outputs = [gr.update(visible=False, value="")] * 10
            
            # Process video cards
            if video_cards:
                for idx, video_info in enumerate(video_cards[:10]):
                    # Create info HTML for each video
                    info_html = f'''
                    <div style="background: white; padding: 16px; border-radius: 12px; border: 1px solid #eee; margin-bottom: 12px; box-shadow: 0 4px 6px rgba(0,0,0,0.02);">
                        <div style="display: flex; justify-content: space-between; align-items: flex-start; margin-bottom: 8px;">
                            <h4 style="margin: 0; color: #1d1d1f; font-weight: 600; font-size: 16px;">Video {idx + 1}</h4>
                            <span style="font-size: 11px; color: #D0021B; background: rgba(208, 2, 27, 0.05); padding: 4px 8px; border-radius: 4px; font-weight: 500;">Score: {video_info['match_score']}</span>
                        </div>
                        <p style="margin: 0 0 12px 0; font-size: 12px; font-family: 'SF Mono', Consolas, monospace; color: #86868b; word-break: break-all;">{video_info['key']}</p>
                        
                        <div style="background: #f5f5f7; padding: 12px; border-radius: 8px; margin-bottom: 12px;">
                            <p style="margin: 0 0 4px 0; font-size: 11px; text-transform: uppercase; letter-spacing: 0.05em; color: #86868b; font-weight: 600;">Summary</p>
                            <p style="margin: 0; font-size: 13px; color: #424245; line-height: 1.4;">{video_info['summary']}</p>
                        </div>
                        
                        <div style="display: flex; gap: 16px; margin: 0;">
                            <div>
                                <p style="margin: 0 0 2px 0; font-size: 11px; color: #86868b;">Duration</p>
                                <p style="margin: 0; font-size: 13px; color: #1d1d1f; font-weight: 500;">{video_info['duration']}s</p>
                            </div>
                            <div>
                                <p style="margin: 0 0 2px 0; font-size: 11px; color: #86868b;">Resolution</p>
                                <p style="margin: 0; font-size: 13px; color: #1d1d1f; font-weight: 500;">{video_info['resolution']}</p>
                            </div>
                        </div>
                    </div>
                    '''
                    
                    video_info_outputs[idx] = gr.update(value=info_html, visible=True)
                    video_player_outputs[idx] = gr.update(value=video_info['path'], visible=True)
            
            # Process document cards
            if document_cards:
                for idx, doc_info in enumerate(document_cards[:10]):
                    # Create document card HTML
                    doc_html = f'''
                    <div style="background: white; padding: 16px; border-radius: 12px; border: 1px solid #eee; margin-bottom: 16px; box-shadow: 0 2px 8px rgba(0,0,0,0.02);">
                        <div style="display: flex; justify-content: space-between; align-items: flex-start; margin-bottom: 10px;">
                            <h4 style="margin: 0; color: #e65100; font-weight: 600; font-size: 16px;">Document {idx + 1}</h4>
                            <span style="font-size: 11px; color: #e65100; background: #fff3e0; padding: 4px 8px; border-radius: 4px; font-weight: 500;">Score: {doc_info['match_score']}</span>
                        </div>
                        <p style="margin: 0 0 12px 0; font-size: 11px; font-family: 'SF Mono', Consolas, monospace; color: #86868b; word-break: break-all;">{doc_info['key']}</p>
                        
                        <div style="background: #fafafa; padding: 12px; border-radius: 8px; margin: 12px 0; border: 1px solid #f5f5f5;">
                            <p style="margin: 0 0 6px 0; font-size: 11px; text-transform: uppercase; letter-spacing: 0.05em; color: #e65100; font-weight: 600;">Summary</p>
                            <p style="margin: 0; font-size: 13px; color: #424245; line-height: 1.5;">{doc_info['summary']}</p>
                        </div>
                        
                        <div style="display: flex; gap: 20px; margin: 10px 0;">
                            <div>
                                <p style="margin: 0 0 2px 0; font-size: 11px; color: #86868b;">Words</p>
                                <p style="margin: 0; font-size: 13px; color: #1d1d1f; font-weight: 500;">{doc_info['word_count']}</p>
                            </div>
                            <div>
                                <p style="margin: 0 0 2px 0; font-size: 11px; color: #86868b;">Sentiment</p>
                                <p style="margin: 0; font-size: 13px; color: #1d1d1f; font-weight: 500;">{doc_info['sentiment']} <span style="color: #86868b;">({doc_info['sentiment_score']})</span></p>
                            </div>
                        </div>
                    '''
                    
                    # Add categories if present
                    if doc_info['categories']:
                        doc_html += '''
                        <div style="background: #e3f2fd; padding: 10px; border-radius: 5px; margin: 10px 0;">
                            <p style="margin: 0 0 5px 0; font-weight: bold; color: #1976d2;"> Categories:</p>
                        '''
                        for category in doc_info['categories']:
                            doc_html += f'<p style="margin: 3px 0; font-size: 12px; color: #555;">• {category}</p>'
                        doc_html += '</div>'
                    
                    # Add entities if present
                    if doc_info['entities']:
                        doc_html += '''
                        <div style="background: #f3e5f5; padding: 10px; border-radius: 5px; margin: 10px 0;">
                            <p style="margin: 0 0 5px 0; font-weight: bold; color: #7b1fa2;"> Named Entities:</p>
                        '''
                        for entity in doc_info['entities']:
                            doc_html += f'<p style="margin: 3px 0; font-size: 12px; color: #555;">• {entity}</p>'
                        doc_html += '</div>'
                    
                    doc_html += '</div>'
                    
                    document_card_outputs[idx] = gr.update(value=doc_html, visible=True)
            
            # Return: summary, table, gallery, then 10 video info, 10 video players, 10 doc cards
            return [summary, table, gallery] + video_info_outputs + video_player_outputs + document_card_outputs
        
        search_btn.click(
            fn=search_and_display_videos,
            inputs=[search_query, search_modality, semantic_threshold, search_mode],
            outputs=[
                search_summary, search_results_table, search_gallery,
                video_info_1, video_info_2, video_info_3, video_info_4, video_info_5,
                video_info_6, video_info_7, video_info_8, video_info_9, video_info_10,
                video_player_1, video_player_2, video_player_3, video_player_4, video_player_5,
                video_player_6, video_player_7, video_player_8, video_player_9, video_player_10,
                doc_card_1, doc_card_2, doc_card_3, doc_card_4, doc_card_5,
                doc_card_6, doc_card_7, doc_card_8, doc_card_9, doc_card_10
            ]
        )
        
        # Copy selected objects
        copy_btn.click(
            fn=copy_selected_objects,
            inputs=[search_results_table, destination_bucket],
            outputs=[operation_status]
        )
        
        # Move selected objects
        move_btn.click(
            fn=move_selected_objects,
            inputs=[search_results_table, destination_bucket],
            outputs=[operation_status]
        )
        
        # Download selected objects
        def download_and_show(table_data):
            status, filepath = download_selected_objects(table_data)
            if filepath:
                return status, gr.update(value=filepath, visible=True)
            else:
                return status, gr.update(visible=False)
        
        download_btn.click(
            fn=download_and_show,
            inputs=[search_results_table],
            outputs=[operation_status, download_file]
        )
        
        # Browse
        list_objects_btn.click(
            fn=list_all_objects,
            inputs=[browse_modality],
            outputs=[dataset_list_html, object_key_dropdown]
        )
        
        # Load object details
        load_object_btn.click(
            fn=load_object_details,
            inputs=[object_key_dropdown],
            outputs=[object_image_preview, object_metadata_display, object_video_preview]
        )
        
        # Video search - populate dropdown
        list_objects_btn.click(
            fn=lambda modality: list_all_objects("video")[1],
            inputs=[browse_modality],
            outputs=[video_search_key]
        )
        video_search_key.change(
            fn=preview_selected_video,
            inputs=[video_search_key],
            outputs=[video_preview_player, video_preview_info]
        )        
        
        # Video frame search with CLIP
        video_search_btn.click(
            fn=search_video_content_semantic,
            inputs=[video_search_key, video_search_query, video_search_threshold],
            outputs=[video_search_results, video_frame_gallery]
        )
    
    return demo


# Launch Application
# ============================================================================
if __name__ == "__main__":
    import os
    import socket
    
    print("=" * 80)
    print("Intelligent Dataset Retrieval Platform - Enterprise Edition")
    print("=" * 80)
    print("\n🔧 Initializing system...")
    
    # Configuration
    TARGET_PORT = 7866
    
    # Auto-detect environment
    def get_server_config():
        """Detect if running locally or on centralized server"""
        hostname = socket.gethostname()
        
        # Try to get the IP address
        try:
            host_ip = socket.gethostbyname(hostname)
        except:
            host_ip = "127.0.0.1"
        
        # Check if running on centralized server
        is_centralized = host_ip.startswith("10.36.97.") or os.getenv("DEPLOY_ENV") == "production"
        
        if is_centralized:
            print(f"🌐 Running on CENTRALIZED SERVER")
            print(f"   Hostname: {hostname}")
            print(f"   Server IP: {host_ip}")
            print(f"   Access URL: http://10.36.97.158:{TARGET_PORT}")
            return {
                "server_name": "0.0.0.0",
                "server_port": TARGET_PORT,
                "share": False,
                "inbrowser": False
            }
        else:
            print(f"💻 Running on LOCAL MACHINE (Mac)")
            print(f"   Hostname: {hostname}")
            print(f"   Local URL: http://127.0.0.1:{TARGET_PORT}")
            print(f"   Network URL: http://{host_ip}:{TARGET_PORT}")
            return {
                "server_name": "0.0.0.0",  # Bind to all interfaces
                "server_port": TARGET_PORT,
                "share": False,
                "inbrowser": True  # Auto-open browser on local
            }
    
    # Initialize analyzers globally
    print("📊 Loading AI models...")
    ai_analyzer = ImageAnalyzer()
    video_analyzer = VideoAnalyzer()
    video_analyzer.set_image_analyzer(ai_analyzer)
    document_analyzer = DocumentAnalyzer()
    
    print("🖥️ Building interface...")
    demo = create_ui()
    
    # Get server configuration
    server_config = get_server_config()
    
    print("\n✅ System ready!")
    print("🌐 Launching web interface...")
    print("=" * 80)
    
    # Launch with detected configuration
    demo.launch(
        server_name=server_config["server_name"],
        server_port=server_config["server_port"],
        share=server_config["share"],
        show_error=True,
        inbrowser=server_config.get("inbrowser", False)
    )
